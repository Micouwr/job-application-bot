"""
Document Parser for Resume Processing
Handles PDF, DOCX, and TXT files with robust error handling
"""

import logging
from pathlib import Path
from typing import Optional

# Mandatory imports - must be included in spec file hiddenimports
try:
    import PyPDF2
    from PyPDF2 import PdfReader
    import docx
    from docx import Document
except ImportError as e:
    # This should never happen in bundled app if spec is correct
    raise RuntimeError(f"Document parser dependencies missing: {e}")

logger = logging.getLogger(__name__)

class DocumentParserError(Exception):
    """Custom exception for document parsing errors"""
    pass

class DocumentParser:
    """Robust document parser with error handling and validation"""

    def __init__(self, max_file_size_mb: int = 10):
        """
        Initialize parser with config

        Args:
            max_file_size_mb: Maximum allowed file size (default 10MB)
        """
        self.max_file_size_mb = max_file_size_mb

    def validate_file(self, file_path: Path) -> None:
        """
        Validate file before parsing

        Raises:
            DocumentParserError: If file is invalid or too large
        """
        if not file_path.exists():
            raise DocumentParserError(f"File not found: {file_path}")

        if not file_path.is_file():
            raise DocumentParserError(f"Not a regular file: {file_path}")

        # Check file size
        size_mb = file_path.stat().st_size / (1024 * 1024)
        if size_mb > self.max_file_size_mb:
            raise DocumentParserError(
                f"File too large: {size_mb:.1f}MB (max {self.max_file_size_mb}MB). "
                "Please use a smaller file or reduce image quality."
            )

        # Check file extension
        valid_extensions = {'.pdf', '.docx', '.doc', '.txt'}
        if file_path.suffix.lower() not in valid_extensions:
            raise DocumentParserError(
                f"Invalid file type: {file_path.suffix}. "
                f"Supported: {', '.join(valid_extensions)}"
            )

    def extract_text(self, file_path: Path) -> str:
        """
        Extract text from any supported document

        Args:
            file_path: Path to PDF, DOCX, or TXT file

        Returns:
            Extracted text as string

        Raises:
            DocumentParserError: If extraction fails
        """
        self.validate_file(file_path)

        try:
            if file_path.suffix.lower() == '.pdf':
                return self._extract_pdf_text(file_path)
            elif file_path.suffix.lower() in ['.docx', '.doc']:
                return self._extract_docx_text(file_path)
            elif file_path.suffix.lower() == '.txt':
                return self._extract_txt_text(file_path)
            else:
                # Should never reach here due to validate_file
                raise DocumentParserError(f"Unsupported file type: {file_path.suffix}")
        except DocumentParserError:
            raise
        except Exception as e:
            logger.exception(f"Unexpected error parsing {file_path}: {e}")
            raise DocumentParserError(f"Failed to parse document: {str(e)}")

    def _extract_pdf_text(self, pdf_path: Path) -> str:
        """
        Extract text from PDF with comprehensive error handling

        Args:
            pdf_path: Path to PDF file

        Returns:
            Extracted text

        Raises:
            DocumentParserError: If PDF is corrupted, encrypted, or image-based
        """
        try:
            reader = PdfReader(pdf_path)

            # Check if PDF is encrypted
            if reader.is_encrypted:
                raise DocumentParserError(
                    "PDF is password-protected. Please remove password protection before uploading."
                )

            text = ""
            page_count = len(reader.pages)

            # Extract text from each page
            for page_num, page in enumerate(reader.pages, 1):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n\n"
                    else:
                        # Page has no extractable text (might be image-based)
                        logger.warning(f"Page {page_num} of {pdf_path} has no extractable text")
                        text += f"[Page {page_num}: No text content (possibly scanned image)]\n\n"
                except Exception as e:
                    logger.warning(f"Failed to extract page {page_num}: {e}")
                    text += f"[Page {page_num}: Extraction failed]\n\n"

            # Validate extracted content
            if len(text.strip()) < 100:
                raise DocumentParserError(
                    "PDF appears to be image-based or corrupted. Text extraction found very little content. "
                    "Consider using OCR software to convert scanned documents to text."
                )

            logger.info(f"Successfully extracted {len(text)} characters from {pdf_path} ({page_count} pages)")
            return text

        except PyPDF2.errors.PdfReadError as e:
            raise DocumentParserError(f"PDF file is corrupted or invalid: {e}")
        except Exception as e:
            logger.exception(f"PDF extraction failed for {pdf_path}")
            raise DocumentParserError(f"Failed to extract PDF: {str(e)}")

    def _extract_docx_text(self, docx_path: Path) -> str:
        """
        Extract text from DOCX with error handling

        Args:
            docx_path: Path to DOCX file

        Returns:
            Extracted text

        Raises:
            DocumentParserError: If DOCX is corrupted
        """
        try:
            doc = Document(docx_path)

            text = ""
            paragraph_count = 0

            # Extract from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n\n"
                    paragraph_count += 1

            # Extract from tables (common in resumes)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + "\t"
                    text += "\n"

            # Validate content
            if len(text.strip()) < 100:
                raise DocumentParserError(
                    "DOCX appears to be empty or corrupted. Extracted very little content."
                )

            logger.info(f"Successfully extracted {len(text)} characters from {docx_path} ({paragraph_count} paragraphs)")
            return text

        except docx.opc.exceptions.PackageNotFoundError:
            raise DocumentParserError(f"DOCX file is corrupted or not a valid Word document: {docx_path}")
        except Exception as e:
            logger.exception(f"DOCX extraction failed for {docx_path}")
            raise DocumentParserError(f"Failed to extract DOCX: {str(e)}")

    def _extract_txt_text(self, txt_path: Path) -> str:
        """
        Extract text from plain text file

        Args:
            txt_path: Path to TXT file

        Returns:
            File content as string

        Raises:
            DocumentParserError: If file can't be read
        """
        try:
            # Try UTF-8 first (most common)
            try:
                text = txt_path.read_text(encoding='utf-8')
            except UnicodeDecodeError:
                # Fall back to system default encoding
                text = txt_path.read_text(encoding='latin-1')

            # Basic validation
            if len(text.strip()) < 50:
                raise DocumentParserError("Text file appears to be empty.")

            logger.info(f"Successfully read {len(text)} characters from {txt_path}")
            return text

        except Exception as e:
            logger.exception(f"Failed to read text file {txt_path}")
            raise DocumentParserError(f"Failed to read text file: {str(e)}")

    def convert_to_txt(self, file_path: Path, output_dir: Optional[Path] = None) -> Path:
        """
        Convert any supported document to TXT file

        Args:
            file_path: Input document path
            output_dir: Directory for output file (default: same as input)

        Returns:
            Path to created TXT file

        Raises:
            DocumentParserError: If conversion fails
        """
        try:
            # Extract text
            text = self.extract_text(file_path)

            # Determine output path
            if output_dir:
                output_dir = Path(output_dir)
                output_dir.mkdir(exist_ok=True)
                output_path = output_dir / f"{file_path.stem}.txt"
            else:
                output_path = file_path.with_suffix('.txt')

            # Write with BOM for Windows compatibility
            output_path.write_text(text, encoding='utf-8-sig')

            logger.info(f"Successfully converted {file_path} to {output_path}")
            return output_path

        except Exception as e:
            logger.exception(f"Conversion failed for {file_path}")
            raise DocumentParserError(f"Conversion failed: {str(e)}")


# Singleton instance for global use
doc_parser = DocumentParser(max_file_size_mb=10)

def parse_document(file_path: Path) -> str:
    """
    Convenience function to parse a document using the singleton instance.
    This is the primary function to be imported by other modules.

    Args:
        file_path: Path to the document.

    Returns:
        Extracted text.

    Raises:
        DocumentParserError: If parsing fails.
    """
    return doc_parser.extract_text(file_path)
