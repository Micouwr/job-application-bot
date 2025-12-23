import tkinter as tk
from tkinter import ttk, messagebox, filedialog, scrolledtext
import ttkthemes
import json
import os
import sys
import threading
import queue
import shutil
from pathlib import Path
from datetime import datetime
from dotenv import load_dotenv
import logging

# Add project root to Python path
project_root = Path(__file__).parent.parent
sys.path.insert(0, str(project_root))
sys.path.insert(0, str(project_root / "utils"))

from config.settings import OUTPUT_PATH, DB_PATH, MIN_MATCH_THRESHOLD
import config.settings
from database import DatabaseManager
from tailor import process_and_tailor_from_gui
from models.resume_model import ResumeModel
from AI.match_analyzer import analyze_match, extract_job_details

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)

class JobAppTkinter:
    def __init__(self, master=None):
        print("MAIN FILE EXECUTED - UNIQUE IDENTIFIER")
        
        # Apply themed style
        self.style = ttkthemes.ThemedStyle(master)
        self.style.set_theme("arc")  # Modern, clean theme
        
        self.master = master
        self.master.title("CareerForge AI - Intelligent Resume Tailoring Tool")
        self.master.geometry("1050x750")
        
        # Set window icon if available
        self._set_window_icon()
        
        # Initialize models and database
        self.db_manager = DatabaseManager()
        self.resume_model = ResumeModel()
        self.selected_resume_path = None
        self.active_resume_id = None
        
        # Ensure output directory exists
        os.makedirs(OUTPUT_PATH, exist_ok=True)
        
        # Create default resume if none exists
        self._ensure_default_resume()
        
        # Queue for thread communication
        self.tailoring_queue = queue.Queue()
        
        # Initialize match analysis variables
        self.match_data = None
        
        # Initialize processing state
        self._processing = False
        
        # Initialize tooltip window
        self._tooltip_window = None
        
        # Initialize current threshold value
        self.current_threshold = MIN_MATCH_THRESHOLD
        
        # Initialize UI
        self._init_ui()
        
        # Check for API key
        self._check_api_key()
        
        # Start queue checker
        self._check_queue()
    
    def _set_window_icon(self):
        """Set the window icon for the application"""
        try:
            # Try different icon formats based on platform
            icon_path = Path(__file__).parent.parent / "assets"
            
            # Try PNG format first for cross-platform compatibility (macOS prefers this)
            png_path = icon_path / "CareerForge_AI.png"
            if png_path.exists() and png_path.stat().st_size > 0:  # Check if file is not empty
                icon_image = tk.PhotoImage(file=str(png_path))
                self.master.iconphoto(True, icon_image)
                return
            
            # Try new CareerForge AI ICO format (Windows)
            careerforge_ico = icon_path / "CareerForge_AI.ico"
            if careerforge_ico.exists():
                self.master.iconbitmap(str(careerforge_ico))
                return
            
            # Try legacy ICO format (Windows)
            ico_path = icon_path / "icon.ico"
            if ico_path.exists():
                self.master.iconbitmap(str(ico_path))
                return
            
            # Fallback to computer.png
            fallback_png = icon_path / "computer.png"
            if fallback_png.exists() and fallback_png.stat().st_size > 0:  # Check if file is not empty
                icon_image = tk.PhotoImage(file=str(fallback_png))
                self.master.iconphoto(True, icon_image)
                return
                
        except Exception as e:
            # Silently fail if icon can't be set - not critical
            pass
    
    def _ensure_default_resume(self):
        """Create default resume if no resumes exist in database"""
        resumes = self.resume_model.list_resumes()
        
        if not resumes:
            default_resume_text = """JOHN DOE
AI Developer & Automation Specialist
Louisville, KY | johndoe@example.com | (555) 123-4567 | linkedin.com/in/johndoe | github.com/johndoe

PROFESSIONAL SUMMARY
Strategic AI Designer and Orchestrator with expertise in architecting intelligent automation solutions and cross-platform applications. Proven track record of designing AI systems that improve efficiency by up to 85%. Strong background in AI strategy, solution architecture, and team leadership utilizing GPT APIs and modern AI platforms.

CORE COMPETENCIES
AI Strategy & Design: Solution architecture, system design, strategic planning
AI Platforms & Tools: GPT-4 API, LangChain, cloud AI services, prompt engineering
Leadership & Collaboration: Team guidance, stakeholder communication, project orchestration
Business Alignment: Requirements translation, ROI optimization, outcome-focused design
Process Optimization: Workflow automation, efficiency improvement, performance metrics
Cross-Platform Solutions: Application design, deployment strategies, user experience

PROFESSIONAL EXPERIENCE

Lead AI Strategist | Tech Innovations Inc. | 2022-Present
- Developed AI-powered resume tailoring system using GPT-4 API, reducing application time by 80%
- Implemented cross-platform desktop application with PyInstaller for 500+ users
- Created automated job application bot that increased interview rate by 3x
- Led team of 3 developers in building machine learning pipeline

Full Stack Developer | Automation Solutions Corp. | 2020-2022
- Architected web scraping automation tools that processed 10,000+ job postings daily
- Designed API integration strategy for LinkedIn, Indeed, and other platforms
- Developed database architecture for candidate tracking and analytics
- Directed cross-platform deployment strategy for Windows, macOS, and Linux

AI Design Specialist | Machine Learning Startup | 2019-2020
- Designed predictive models architecture for applicant tracking systems (ATS)
- Created NLP solution framework for resume optimization
- Built custom algorithm architecture for job matching and candidate ranking

EDUCATION
Bachelor of Science in Computer Science
University of California, Berkeley | 2019
Relevant Coursework: Machine Learning, AI Systems, Data Structures, Algorithms

KEY PROJECTS
CareerForge AI - AI Resume Tailorer
- Designed intelligent resume tailoring system using GPT-4 API
- Architected cross-platform solution for Windows, macOS, and Linux
- Orchestrated development process with technical team

Enterprise Automation Suite
- Designed automation architecture for job search and application tracking
- Directed integration with 15+ job boards via API and web services
- Managed cross-platform deployment strategy

AI Performance Optimizer
- Designed machine learning model architecture for resume analysis
- Created optimization framework for ATS performance and keyword targeting
- Guided implementation with measurable 90% user satisfaction

PROFESSIONAL DEVELOPMENT
- Continuous learning in AI strategy and emerging technologies
- Active participant in AI leadership and design communities
- Member: AI Strategy Association, Technology Leadership Network

REFERENCES
Available upon request. Portfolio of designed solutions and architectural diagrams accessible via repository.
"""
            
            # Save to file
            default_path = OUTPUT_PATH / "default_resume.txt"
            with open(default_path, 'w', encoding='utf-8') as f:
                f.write(default_resume_text)
            
            # Add to database
            self.resume_model.add_resume(str(default_path), "Default Resume", is_active=True)
            logging.info("Default resume created successfully")
    
    def _init_ui(self):
        """Initialize the main UI components"""
        # Configure styles
        self.style.configure('TNotebook.Tab', font=('Arial', 10, 'bold'), padding=[10, 5])
        self.style.configure('TButton', font=('Arial', 9), padding=[5, 2])
        self.style.configure('TLabel', font=('Arial', 10))
        self.style.configure('Bold.TLabel', font=('Arial', 10, 'bold'))
        self.style.configure('Title.TLabel', font=('Arial', 14, 'bold'))
        
        # Main container
        main_frame = ttk.Frame(self.master, padding="15")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        self.master.columnconfigure(0, weight=1)
        self.master.rowconfigure(0, weight=1)
        
        # Create notebook for tabs
        self.notebook = ttk.Notebook(main_frame)
        self.notebook.grid(row=0, column=0, columnspan=4, sticky=(tk.W, tk.E, tk.N, tk.S), pady=(0, 10))
        
        # Create tabs in desired order
        self._create_add_job_tab()  # Job Management
        self._create_import_tab()   # Import from LinkedIn/Email
        self._create_resume_mgmt_tab()
        self._create_tailored_docs_tab()
        self._create_custom_prompt_tab()
        self._create_settings_tab()  # Settings/Preferences
        self._create_output_tab()  # OUTPUT & LOGS moved to final position
        
        # Make notebook expandable
        main_frame.columnconfigure(0, weight=1)
        main_frame.rowconfigure(0, weight=1)
    
    def _create_add_job_tab(self):
        """Create the Add Job Application tab"""
        tab = ttk.Frame(self.notebook)
        self.notebook.add(tab, text="Job Management")
        
        # Title
        title_label = ttk.Label(tab, text="Job Application Details", style='Title.TLabel')
        title_label.grid(row=0, column=0, columnspan=3, pady=(0, 15), sticky=tk.W)
        
        # Job Details Section
        ttk.Label(tab, text="Job Title:", style='Bold.TLabel').grid(row=1, column=0, sticky=tk.W, pady=5)
        self.job_title_entry = ttk.Entry(tab, width=50, font=('Arial', 10))
        self.job_title_entry.grid(row=1, column=1, columnspan=2, sticky=(tk.W, tk.E), pady=5)
        
        ttk.Label(tab, text="Company:", style='Bold.TLabel').grid(row=2, column=0, sticky=tk.W, pady=5)
        self.company_entry = ttk.Entry(tab, width=50, font=('Arial', 10))
        self.company_entry.grid(row=2, column=1, columnspan=2, sticky=(tk.W, tk.E), pady=5)
        
        # Role Level Selection
        ttk.Label(tab, text="Role Level:", style='Bold.TLabel').grid(row=3, column=0, sticky=tk.W, pady=5)
        self.role_var = tk.StringVar(value="Standard")
        role_combo = ttk.Combobox(tab, textvariable=self.role_var, values=["Standard", "Senior", "Lead", "Principal"], 
                                state='readonly', font=('Arial', 10), width=47)
        role_combo.grid(row=3, column=1, columnspan=2, sticky=(tk.W, tk.E), pady=5)
        
        # Add tooltip with enhanced role definitions
        role_tooltip = "ROLE LEVEL DEFINITIONS & USAGE:\n\n"
        role_tooltip += "STANDARD (0-5 years experience):\n"
        role_tooltip += "* Focus on core technical skills and direct contributions\n"
        role_tooltip += "* Primary responsibilities include hands-on work and task execution\n"
        role_tooltip += "* Limited supervision with clear direction from managers\n\n"
        role_tooltip += "SENIOR (5-10 years experience):\n"
        role_tooltip += "* Demonstrated expertise in core technologies and methodologies\n"
        role_tooltip += "* Mentors junior team members and provides technical guidance\n"
        role_tooltip += "* Contributes to architectural decisions and project planning\n"
        role_tooltip += "* Works independently with minimal supervision\n\n"
        role_tooltip += "LEAD (8-15 years experience):\n"
        role_tooltip += "* Manages teams or significant project components\n"
        role_tooltip += "* Responsible for resource allocation and timeline management\n"
        role_tooltip += "* Interfaces with stakeholders and translates business needs\n"
        role_tooltip += "* Drives process improvements and best practices\n\n"
        role_tooltip += "PRINCIPAL (12+ years experience):\n"
        role_tooltip += "* Recognized expert and thought leader in the field\n"
        role_tooltip += "* Sets technical direction for large-scale initiatives\n"
        role_tooltip += "* Influences organizational strategy and innovation\n"
        role_tooltip += "* Represents company externally at conferences and industry events\n\n"
        role_tooltip += "HOW TO CHOOSE ROLE LEVEL:\n"
        role_tooltip += "* Match the role level to the job posting requirements\n"
        role_tooltip += "* When in doubt, start with STANDARD and adjust based on match scores\n"
        role_tooltip += "* Higher role levels require demonstrated leadership experience\n\n"
        role_tooltip += "RELATIONSHIP TO CUSTOM PROMPTS:\n"
        role_tooltip += "* Role level determines base template used for tailoring\n"
        role_tooltip += "* Custom prompts are applied IN ADDITION to role level adjustments\n"
        role_tooltip += "* Custom prompts can override or enhance role-level behavior\n"
        role_tooltip += "* Both work together: Role level + Custom prompt = Final output\n\n"
        role_tooltip += "BEST PRACTICES:\n"
        role_tooltip += "* Use role level for general experience level matching\n"
        role_tooltip += "* Use custom prompts for specific company/role requirements\n"
        role_tooltip += "* Test both together to ensure desired results"
        role_combo.tooltip = role_tooltip
        
        # Bind hover event to show tooltip
        def show_role_tooltip(event):
            self.master.after(100, lambda: self._show_tooltip(event, role_tooltip))
        
        role_combo.bind('<Enter>', show_role_tooltip)
        role_combo.bind('<Leave>', lambda e: self._hide_tooltip())
        
        # Role Level Help Text
        role_help = ttk.Label(tab, text="Select role level that matches the job posting (see README for guidance)", 
                             font=('Arial', 9), foreground='blue')
        role_help.grid(row=4, column=1, columnspan=2, sticky=tk.W, pady=(0, 10))
        
        ttk.Label(tab, text="Job Description:", style='Bold.TLabel').grid(row=5, column=0, sticky=tk.W, pady=5)
        self.job_desc_text = scrolledtext.ScrolledText(tab, width=80, height=15, wrap=tk.WORD, font=('Arial', 10))
        self.job_desc_text.grid(row=5, column=1, columnspan=2, sticky=(tk.W, tk.E, tk.N, tk.S), pady=5)
        
        ttk.Label(tab, text="Job URL:", style='Bold.TLabel').grid(row=6, column=0, sticky=tk.W, pady=5)
        self.job_url_entry = ttk.Entry(tab, width=50, font=('Arial', 10))
        self.job_url_entry.grid(row=6, column=1, columnspan=2, sticky=(tk.W, tk.E), pady=5)
        
        # Status section
        status_frame = ttk.LabelFrame(tab, text="Status", padding="10")
        status_frame.grid(row=7, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=10)
        status_frame.columnconfigure(0, weight=1)
        
        # Status label for match analysis feedback
        self.status_label = ttk.Label(status_frame, text="Ready", font=('Arial', 9), foreground='green')
        self.status_label.grid(row=0, column=0, sticky=tk.W, pady=2)
        
        # Match score display
        self.match_label = ttk.Label(status_frame, text="Match Score: Not analyzed", style='Bold.TLabel', foreground='blue')
        self.match_label.grid(row=1, column=0, sticky=tk.W, pady=2)
        
        # Buttons - Reordered as requested
        button_frame = ttk.Frame(tab)
        button_frame.grid(row=8, column=0, columnspan=4, pady=15)
        
        # Add Analyze Match button (first)
        self.analyze_button = ttk.Button(button_frame, text="Analyze Match", command=self.analyze_match)
        self.analyze_button.grid(row=0, column=0, padx=5)
        
        # Start Tailoring button (second)
        self.start_button = ttk.Button(button_frame, text="Start Tailoring", command=self.start_tailoring)
        self.start_button.grid(row=0, column=1, padx=5)
        
        # Clear Fields button (third)
        self.clear_button = ttk.Button(button_frame, text="Clear Fields", command=self.clear_fields)
        self.clear_button.grid(row=0, column=2, padx=5)
        
        # Quit button (fourth)
        self.quit_button = ttk.Button(button_frame, text="Quit", command=self.quit_application)
        self.quit_button.grid(row=0, column=3, padx=5)
        
        # Configure grid weights
        tab.columnconfigure(1, weight=1)
        tab.rowconfigure(5, weight=1)
    
    def _create_resume_mgmt_tab(self):
        """Create the Resume Management tab"""
        tab = ttk.Frame(self.notebook)
        self.notebook.add(tab, text="Resume Management")
        
        # Title
        title_label = ttk.Label(tab, text="Manage Your Resumes", style='Title.TLabel')
        title_label.grid(row=0, column=0, columnspan=4, pady=(0, 15), sticky=tk.W)
        
        # Resume List Section
        ttk.Label(tab, text="Available Resumes:", style='Bold.TLabel').grid(row=1, column=0, sticky=tk.W, pady=5)
        
        # Create treeview for resumes
        columns = ('Name', 'Path', 'Active')
        self.resume_tree = ttk.Treeview(tab, columns=columns, show='headings', height=10)
        self.resume_tree.heading('Name', text='Resume Name')
        self.resume_tree.heading('Path', text='File Path')
        self.resume_tree.heading('Active', text='Active')
        
        self.resume_tree.column('Name', width=150)
        self.resume_tree.column('Path', width=400)
        self.resume_tree.column('Active', width=50)
        
        self.resume_tree.grid(row=2, column=0, columnspan=4, sticky=(tk.W, tk.E, tk.N, tk.S), pady=5)
        
        # Add scrollbar
        scrollbar = ttk.Scrollbar(tab, orient=tk.VERTICAL, command=self.resume_tree.yview)
        self.resume_tree.configure(yscrollcommand=scrollbar.set)
        scrollbar.grid(row=2, column=3, sticky=(tk.N, tk.S))
        
        # Bind selection event
        self.resume_tree.bind('<<TreeviewSelect>>', self._on_resume_select)
        
        # Resume Preview Section
        ttk.Label(tab, text="Resume Preview:", style='Bold.TLabel').grid(row=3, column=0, sticky=tk.W, pady=(15, 5))
        self.resume_preview = scrolledtext.ScrolledText(tab, width=80, height=10, wrap=tk.WORD, font=('Arial', 10))
        # Configure the text widget to have better visual appearance
        self.resume_preview.config(padx=15, pady=15)
        self.resume_preview.grid(row=4, column=0, columnspan=4, sticky=(tk.W, tk.E, tk.N, tk.S), pady=5)
        
        # Buttons
        button_frame = ttk.Frame(tab)
        button_frame.grid(row=5, column=0, columnspan=4, pady=15)
        
        self.upload_button = ttk.Button(button_frame, text="Upload Resume", command=self.upload_resume)
        self.upload_button.grid(row=0, column=0, padx=5)
        
        self.delete_button = ttk.Button(button_frame, text="Delete Selected", command=self.delete_selected_resume)
        self.delete_button.grid(row=0, column=1, padx=5)
        
        self.set_active_button = ttk.Button(button_frame, text="Set as Active", command=self.set_active_resume)
        self.set_active_button.grid(row=0, column=2, padx=5)
        
        # Refresh resume list
        self._refresh_resume_list()
        
        # Add quit button
        quit_button = ttk.Button(button_frame, text="Quit", command=self.quit_application)
        quit_button.grid(row=0, column=3, padx=5)
        
        # Configure grid weights
        tab.columnconfigure(0, weight=1)
        tab.rowconfigure(2, weight=1)
        tab.rowconfigure(4, weight=1)
    
    def _create_settings_tab(self):
        """Create the Settings/Preferences tab"""
        tab = ttk.Frame(self.notebook)
        self.notebook.add(tab, text="Settings")
        
        # Title
        title_label = ttk.Label(tab, text="Application Settings", style='Title.TLabel')
        title_label.grid(row=0, column=0, sticky=tk.W, pady=(0, 15))
        
        # Minimum Match Threshold Setting
        threshold_frame = ttk.LabelFrame(tab, text="Minimum Match Threshold", padding="10")
        threshold_frame.grid(row=1, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=10)
        threshold_frame.columnconfigure(1, weight=1)
        
        ttk.Label(threshold_frame, text="Current Threshold:", style='Bold.TLabel').grid(row=0, column=0, sticky=tk.W, pady=5)
        self.threshold_value_label = ttk.Label(threshold_frame, text=f"{self.current_threshold}%", font=('Arial', 10, 'bold'), foreground='blue')
        self.threshold_value_label.grid(row=0, column=1, sticky=tk.W, pady=5, padx=(10, 0))
        
        ttk.Label(threshold_frame, text="Adjust Threshold:", style='Bold.TLabel').grid(row=1, column=0, sticky=tk.W, pady=5)
        
        # Slider for threshold adjustment
        self.threshold_slider = ttk.Scale(threshold_frame, from_=50, to=95, orient=tk.HORIZONTAL, length=300, command=self._on_threshold_slider_change)
        self.threshold_slider.set(self.current_threshold)
        self.threshold_slider.grid(row=1, column=1, sticky=(tk.W, tk.E), pady=5, padx=(10, 0))
        
        # Entry for manual threshold input
        ttk.Label(threshold_frame, text="Or enter manually:", font=('Arial', 9)).grid(row=2, column=0, sticky=tk.W, pady=5)
        self.threshold_entry = ttk.Entry(threshold_frame, width=10, font=('Arial', 10))
        self.threshold_entry.insert(0, str(self.current_threshold))
        self.threshold_entry.grid(row=2, column=1, sticky=tk.W, pady=5, padx=(10, 0))
        
        # Apply button
        apply_button = ttk.Button(threshold_frame, text="Apply Changes", command=self._apply_threshold_change)
        apply_button.grid(row=3, column=1, sticky=tk.W, pady=10, padx=(10, 0))
        
        # Reset to default button
        reset_button = ttk.Button(threshold_frame, text="Reset to Default (80%)", command=self._reset_threshold_to_default)
        reset_button.grid(row=3, column=1, sticky=tk.W, pady=10, padx=(120, 0))
        
        # Enhanced information text with comprehensive guidance
        info_text = """
How Match Threshold Works:
• The match threshold determines when the "Start Tailoring" button becomes available
• Match Score ≥ Threshold: Tailoring is enabled
• Match Score < Threshold: Tailoring is disabled to prevent poor-quality results

Threshold Strategy Guidelines:
• 70-75%: Aggressive applications (lower match tolerance, more opportunities)
• 80%: Balanced approach (recommended default for most users)
• 85-90%: Conservative applications (higher match requirements, fewer but better-targeted applications)

Best Practices:
1. Start with the default 80% threshold
2. Adjust based on your job application results
3. Lower thresholds for broader opportunities
4. Higher thresholds for targeted, high-match applications
5. Monitor detailed match analysis to understand score factors

Using This Interface:
• Drag the slider to adjust threshold from 50% to 95%
• Or enter a specific value in the text box
• Click "Apply Changes" to save your settings
• Use "Reset to Default (80%)" to restore recommended setting
        """
        info_label = ttk.Label(threshold_frame, text=info_text.strip(), font=('Arial', 9), foreground='blue')
        info_label.grid(row=4, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=10)
        
        # Add tooltip bindings for interactive help
        self.threshold_slider.bind("<Enter>", lambda e: self._show_tooltip(e, "Drag to adjust match threshold from 50% to 95%"))
        self.threshold_slider.bind("<Leave>", lambda e: self._hide_tooltip())
        self.threshold_entry.bind("<Enter>", lambda e: self._show_tooltip(e, "Enter a specific threshold value (50-95%)"))
        self.threshold_entry.bind("<Leave>", lambda e: self._hide_tooltip())
        apply_button.bind("<Enter>", lambda e: self._show_tooltip(e, "Save your threshold settings"))
        apply_button.bind("<Leave>", lambda e: self._hide_tooltip())
        reset_button.bind("<Enter>", lambda e: self._show_tooltip(e, "Restore default threshold (80%)"))
        reset_button.bind("<Leave>", lambda e: self._hide_tooltip())
        
        # Add help button for additional information
        help_button = ttk.Button(tab, text="?", width=3, command=self._show_threshold_help)
        help_button.grid(row=0, column=1, sticky=tk.E, padx=10)
        
        # Configure grid weights
        tab.columnconfigure(0, weight=1)
    
    def _show_threshold_help(self):
        """Show detailed help dialog for threshold configuration"""
        help_text = """
COMPREHENSIVE GUIDE: MATCH THRESHOLD CONFIGURATION

WHAT IS THE MATCH THRESHOLD?
The match threshold is a percentage that determines how closely your resume must match a job description to enable tailoring.

HOW IT WORKS:
• When you analyze a job description, the system calculates a match score
• If the score is equal to or greater than your threshold, tailoring is enabled
• If the score is below your threshold, tailoring is disabled to prevent poor results

STRATEGIC THRESHOLD SETTINGS:

AGGRESSIVE (70-75%)
• Enables tailoring for more job opportunities
• May produce less targeted results
• Best for: High-volume job searching, entry-level positions

BALANCED (80% - DEFAULT)
• Good balance of opportunity and targeting
• Recommended for most users
• Best for: General job searching across experience levels

CONSERVATIVE (85-90%)
• Only enables tailoring for high-match opportunities
• Produces more targeted results
• Best for: Targeted job searching, senior-level positions

INTERFACE CONTROLS:

SLIDER CONTROL:
• Drag left/right to adjust from 50% to 95%
• Real-time display of current value
• Smooth, intuitive adjustment

MANUAL ENTRY:
• Type exact percentage value
• Valid range: 50-95%
• Automatic validation

APPLY CHANGES:
• Saves your threshold setting
• Applies immediately to current and future analyses
• Confirmation message with reminder tips

RESET TO DEFAULT:
• Restores 80% threshold
• One-click convenience
• Helpful if you're unsure of optimal settings

BEST PRACTICES:
1. START WITH DEFAULT (80%)
2. ADJUST BASED ON RESULTS
3. LOWER FOR MORE OPPORTUNITIES
4. RAISE FOR BETTER TARGETING
5. REVIEW MATCH ANALYSES
        """
        
        # Create help window
        help_window = tk.Toplevel(self.master)
        help_window.title("Match Threshold Help")
        help_window.geometry("600x700")
        help_window.resizable(True, True)
        
        # Center the window
        help_window.transient(self.master)
        help_window.grab_set()
        
        # Create scrolled text widget for help content
        help_text_widget = scrolledtext.ScrolledText(
            help_window, 
            width=70, 
            height=40,
            wrap=tk.WORD,
            font=('Arial', 10),
            padx=10, 
            pady=10
        )
        help_text_widget.insert('1.0', help_text.strip())
        help_text_widget.config(state='disabled')  # Make read-only
        help_text_widget.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Add close button
        close_button = ttk.Button(help_window, text="Close", command=help_window.destroy)
        close_button.pack(pady=10)
    
    def _on_threshold_slider_change(self, value):
        """Handle slider value changes"""
        threshold = int(float(value))
        self.threshold_value_label.config(text=f"{threshold}%")
        # Only update entry if it exists (prevents error during initialization)
        if hasattr(self, 'threshold_entry') and self.threshold_entry:
            self.threshold_entry.delete(0, tk.END)
            self.threshold_entry.insert(0, str(threshold))
    
    def _apply_threshold_change(self):
        """Apply the new threshold value"""
        try:
            new_threshold = int(self.threshold_entry.get())
            if new_threshold < 50 or new_threshold > 95:
                messagebox.showerror("Invalid Threshold", "Threshold must be between 50 and 95 percent.")
                return
            
            self.current_threshold = new_threshold
            self.threshold_slider.set(new_threshold)
            self.threshold_value_label.config(text=f"{new_threshold}%")
            
            messagebox.showinfo("Success", f"Minimum match threshold updated to {new_threshold}%\n\nRemember: Lower thresholds (70-75%) = More opportunities\nHigher thresholds (85-90%) = Better-targeted applications")
            self._log_message(f"Minimum match threshold updated to {new_threshold}%", "info")
            
        except ValueError:
            messagebox.showerror("Invalid Input", "Please enter a valid number for the threshold.")
    
    def _reset_threshold_to_default(self):
        """Reset threshold to default value (80%)"""
        default_threshold = 80
        self.current_threshold = default_threshold
        self.threshold_slider.set(default_threshold)
        self.threshold_entry.delete(0, tk.END)
        self.threshold_entry.insert(0, str(default_threshold))
        self.threshold_value_label.config(text=f"{default_threshold}%")
        
        messagebox.showinfo("Success", f"Minimum match threshold reset to default ({default_threshold}%)")
        self._log_message(f"Minimum match threshold reset to default ({default_threshold}%)", "info")
    
    def _create_output_tab(self):
        """Create the Output/Logs tab"""
        tab = ttk.Frame(self.notebook)
        self.notebook.add(tab, text="Output & Logs")
        
        # Title
        title_label = ttk.Label(tab, text="Application Logs", style='Title.TLabel')
        title_label.grid(row=0, column=0, sticky=tk.W, pady=(0, 15))
        
        ttk.Label(tab, text="Application Logs:", style='Bold.TLabel').grid(row=1, column=0, sticky=tk.W, pady=5)
        
        self.log_text = scrolledtext.ScrolledText(tab, width=80, height=25, wrap=tk.WORD, font=('Arial', 10))
        self.log_text.grid(row=2, column=0, columnspan=2, sticky=(tk.W, tk.E, tk.N, tk.S), pady=5)
        
        # Add quit button
        quit_button = ttk.Button(tab, text="Quit", command=self.quit_application)
        quit_button.grid(row=3, column=0, pady=15, sticky=tk.W)
        
        # Configure grid weights
        tab.columnconfigure(0, weight=1)
        tab.rowconfigure(2, weight=1)
    
    def _on_resume_select(self, event):
        """Handle resume selection in treeview"""
        selection = self.resume_tree.selection()
        if selection:
            item = self.resume_tree.item(selection[0])
            resume_path = item['values'][1]
            
            try:
                with open(resume_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                    self.resume_preview.delete('1.0', tk.END)
                    
                    # Add content with improved visual formatting
                    lines = content.split('\n')
                    for i, line in enumerate(lines):
                        # Identify and highlight section headers
                        is_header = line.strip().upper() == line.strip() and len(line.strip()) < 50 and not line.strip().endswith('.') and len(line.strip()) > 0
                        
                        self.resume_preview.insert(tk.END, line + '\n')
                        if is_header:
                            # Tag headers to make them stand out
                            start_pos = self.resume_preview.index(f"{i+1}.0")
                            end_pos = self.resume_preview.index(f"{i+1}.end")
                            self.resume_preview.tag_add('header', start_pos, end_pos)
                    
                    # Configure header tags to make them stand out
                    self.resume_preview.tag_config('header', font=('Arial', 11, 'bold'), foreground='blue', spacing1=8, spacing3=8)
                    
                    # Add padding around the content
                    self.resume_preview.config(padx=20, pady=20)
            except Exception as e:
                self._log_message(f"Error loading resume preview: {e}", "error")
    
    def upload_resume(self):
        """Upload a new resume file"""
        file_path = filedialog.askopenfilename(
            title="Select Resume File",
            filetypes=[
                ("PDF files", "*.pdf"),
                ("Text files", "*.txt"),
                ("Word files", "*.docx"),
                ("All files", "*.*")
            ]
        )
        
        if not file_path:
            return
        
        try:
            # Get file name
            name = Path(file_path).stem
            
            # Handle PDF files
            if file_path.lower().endswith('.pdf'):
                try:
                    import PyPDF2
                    with open(file_path, 'rb') as pdf_file:
                        pdf_reader = PyPDF2.PdfReader(pdf_file)
                        text_content = ""
                        for page in pdf_reader.pages:
                            text_content += page.extract_text() + "\n"
                    
                    # Normalize text to fix common PDF extraction issues
                    # Remove excessive line breaks and fix word splits
                    import re
                    # First, fix common word splits across lines
                    text_content = re.sub(r'([a-zA-Z])-\n([a-zA-Z])', r'\1\2', text_content)  # Join hyphenated words
                    # Fix names that might be split into individual letters
                    text_content = re.sub(r'([A-Z])\n([A-Z])\n([A-Z])\n([A-Z])\n([A-Z])\n([A-Z])\n([A-Z])', r'\1\2\3\4\5\6\7', text_content)  # Fix 7-letter names
                    text_content = re.sub(r'([A-Z])\n([A-Z])\n([A-Z])\n([A-Z])\n([A-Z])\n([A-Z])', r'\1\2\3\4\5\6', text_content)  # Fix 6-letter names like WILLIAM
                    text_content = re.sub(r'\n([A-Z])\n([A-Z])\n([A-Z])\n([A-Z])\n([A-Z])\n', r'\1\2\3\4\5\n', text_content)  # Fix 5-letter names
                    text_content = re.sub(r'\n([A-Z])\n([A-Z])\n([A-Z])\n([A-Z])\n', r'\1\2\3\4\n', text_content)  # Fix 4-letter names
                    text_content = re.sub(r'\n([A-Z])\n([A-Z])\n([A-Z])\n', r'\1\2\3\n', text_content)  # Fix 3-letter names
                    text_content = re.sub(r'\n([A-Z])\n([A-Z])\n', r'\1\2\n', text_content)  # Fix 2-letter names
                    
                    # Also fix potential spaces inserted in the middle of names and URLs
                    text_content = re.sub(r'(WILL)\s+(IAM)', r'WILLIAM', text_content)  # Fix WILL I AM
                    text_content = re.sub(r'(WILLI)\s+(AM)', r'WILLIAM', text_content)  # Fix WILLI AM
                    text_content = re.sub(r'(RYA)\s+(N)', r'RYAN', text_content)  # Fix RYA N
                    # Fix LinkedIn URL that might be split
                    text_content = re.sub(r'linkedin\.com/in/r\s+yanmicou', r'linkedin.com/in/ryanmicou', text_content)  # Fix split LinkedIn URL
                    text_content = re.sub(r'linkedin\.com/in/(\w)\s+(\w+)', r'linkedin.com/in/\1\2', text_content)  # Fix general LinkedIn URL splits
                    
                    # Improve resume formatting by intelligently handling line breaks
                    # First, handle phone numbers that might be split
                    text_content = re.sub(r'(\(\d{3}\))\s*\n\s*(\d{3})', r'\1 \2', text_content)  # Fix phone: (502)\n777 -> (502) 777
                    text_content = re.sub(r'(\d{3})\s*\n\s*(\d{3})\s*\n\s*(\d{4})', r'\1-\2-\3', text_content)  # Fix phone: 502\n777\n7526 -> 502-777-7526
                    
                    # Handle email addresses that might be split
                    text_content = re.sub(r'(\w+)\s*\n\s*(@)\s*\n\s*(\w+)', r'\1\2\3', text_content)  # Fix email splits
                    text_content = re.sub(r'(\w+@\w+)\s*\n\s*(\w+\.\w+)', r'\1.\2', text_content)  # Fix email domain splits
                    
                    # Improved approach: First do all word split fixes, then reconstruct paragraphs
                    # This ensures that fixes like 'cust om' -> 'custom' happen before paragraph reconstruction
                    
                    # Fix specific known word splits that are common in PDF extraction
                    text_content = re.sub(r'Oper\s+ations', r'Operations', text_content)  # Fix "Oper ations"
                    text_content = re.sub(r'compr\s+ehensiv\s+e', r'comprehensive', text_content)  # Fix "compr ehensiv e"
                    text_content = re.sub(r'envir\s+onments', r'environments', text_content)  # Fix "envir onments"
                    text_content = re.sub(r'deliv\s+ery', r'delivery', text_content)  # Fix "deliv ery"
                    text_content = re.sub(r'procedur\s+es', r'procedures', text_content)  # Fix "procedur es"
                    text_content = re.sub(r'work\s+ﬂow', r'workflow', text_content)  # Fix "workﬂow" with special character
                    text_content = re.sub(r'conﬁgurations', r'configurations', text_content)  # Fix special character 'ﬁ'
                    text_content = re.sub(r'certiﬁed', r'certified', text_content)  # Fix special character 'ﬁ'
                    text_content = re.sub(r'certiﬁcations', r'certifications', text_content)  # Fix special character 'ﬁ'
                    
                    # Fix common concatenated words that appear in PDF extraction
                    text_content = re.sub(r'Manage([a-z]+)team', r'Manage \1 team', text_content)  # Fix "Managedateam"
                    text_content = re.sub(r'aregulated', r'a regulated', text_content)  # Fix "aregulated"
                    text_content = re.sub(r'acentralized', r'a centralized', text_content)  # Fix "acentralized"
                    
                    # Fix other common concatenated words
                    text_content = re.sub(r'deploy([a-z]+)om', r'deploy \1 om', text_content)  # Fix "deploycust om" -> "deploy cust om"
                    
                    # Fix common patterns where spaces were incorrectly inserted in the middle of words
                    text_content = re.sub(r'(\w{2,})-\s+(\w{2,})', r'\1-\2', text_content)  # Fix "high-v olume" -> "high-volume"
                    text_content = re.sub(r'(\w{2,})\s+([a-z])\s+(\w{2,})', r'\1\2\3', text_content)  # Fix "oper ations" -> "operations"
                    
                    # Fix specific issues from current output
                    text_content = re.sub(r'cust\s+om', r'custom', text_content)  # Fix "cust om" -> "custom"
                    text_content = re.sub(r'Manage\s+da', r'Managed a', text_content)  # Fix "Manage da" -> "Managed a"
                    
                    # Handle general case of single character splits within words
                    text_content = re.sub(r'(\w{3,})\s+(\w)\s+(\w{3,})', r'\1\2\3', text_content)  # Fix word-char-word pattern like "compr e hensive"
                    
                    # Fix any remaining two-word concatenations that should be separated
                    text_content = re.sub(r'([a-z])([A-Z])', r'\1 \2', text_content)  # Insert space between lowercase and uppercase: 'deliveryProven' -> 'delivery Proven'
                    
                    # Final cleanup: fix any remaining obvious word splits that we can identify
                    text_content = text_content.replace('aregulated', 'a regulated')  # Fix "aregulated"
                    text_content = text_content.replace('acentralized', 'a centralized')  # Fix "acentralized"
                    text_content = text_content.replace('deploycust om', 'deploy cust om')  # Fix "deploycust om"
                    text_content = text_content.replace('Saa S', 'SaaS')  # Fix "Saa S"
                    text_content = text_content.replace('Compu Com', 'CompuCom')  # Fix "Compu Com"
                    text_content = text_content.replace('Accu Code', 'AccuCode')  # Fix "Accu Code"
                    text_content = text_content.replace('Code Louisville', 'CodeLouisville')  # Fix "Code Louisville"
                    text_content = text_content.replace('Comp TIA', 'CompTIA')  # Fix "Comp TIA"
                    
                    # Additional fixes for remaining word splits
                    text_content = re.sub(r'infrastructur\s+es', r'infrastructure', text_content)
                    text_content = re.sub(r'high-v\s+olume', r'high-volume', text_content)
                    text_content = re.sub(r'oper\s+ations', r'operations', text_content)
                    text_content = re.sub(r'oper\s+ational', r'operational', text_content)
                    text_content = re.sub(r'high-le\s+verage', r'high-leverage', text_content)
                    text_content = re.sub(r'oppor\s+tunities', r'opportunities', text_content)
                    text_content = re.sub(r'deliv\s+ering', r'delivering', text_content)
                    text_content = re.sub(r'time-t\s+o-mill', r'time-to-mill', text_content)
                    text_content = re.sub(r'workﬂow', r'workflow', text_content)  # Fix special character
                    text_content = re.sub(r'Py\s+Installer', r'PyInstaller', text_content)
                    text_content = re.sub(r'ina\s+regulated', r'in a regulated', text_content)
                    text_content = re.sub(r'Createda', r'Created a', text_content)
                    
                    # Fix additional word splits seen in the current output
                    text_content = re.sub(r'exper\s+tise', r'expertise', text_content)  # Fix "exper tise"
                    text_content = re.sub(r'CAP\s+ABILI\s+TIES', r'CAPABILITIES', text_content)  # Fix "CAP ABILI TIES"
                    text_content = re.sub(r'PROJEC\s+TS', r'PROJECTS', text_content)  # Fix "PROJEC TS"
                    text_content = re.sub(r'certiﬁ\s+cations', r'certifications', text_content)  # Fix "certiﬁ cations" with special character
                    text_content = re.sub(r'Governance\s+Pro\s+jects', r'Governance Projects', text_content)  # Fix "Governance Projects" splits
                    text_content = re.sub(r'\s+\s+', ' ', text_content)  # Replace multiple spaces with single space
                    
                    # Now do the paragraph reconstruction after all word fixes
                    lines = text_content.split('\n')
                    processed_lines = []
                    i = 0
                    
                    # Define resume section headers that should be on separate lines
                    section_headers = ['summary', 'experience', 'education', 'skills', 'projects', 'certifications', 
                                     'professional summary', 'work experience', 'professional experience',
                                     'technical skills', 'core capabilities', 'ai projects',
                                     'professional experience', 'education & certifications']
                    
                    while i < len(lines):
                        current_line = lines[i].strip()
                        if not current_line:
                            i += 1
                            continue
                        
                        # Check if this line is a section header
                        is_section_header = any(header in current_line.lower() for header in section_headers) or \
                                          current_line.isupper() and len(current_line) <= 50  # Likely a section header if all caps and short
                        
                        # Check if this line is a job entry (contains years)
                        is_job_entry = any(char.isdigit() for char in current_line) and ('–' in current_line or '-' in current_line)
                        
                        # Check if this line starts with special characters (bullets, etc.)
                        is_list_item = current_line.startswith(('●', '○', '§', '•', '-', '—', '|'))
                        
                        if is_section_header or is_job_entry or is_list_item:
                            # These should remain on separate lines
                            processed_lines.append(current_line)
                            i += 1
                        else:
                            # Join this line and subsequent lines that don't look like headers
                            paragraph = current_line
                            i += 1
                            
                            while i < len(lines):
                                next_line = lines[i].strip()
                                
                                if not next_line:
                                    i += 1
                                    continue  # Skip blank lines but keep joining
                                
                                # Check if next line should start a new paragraph
                                next_is_header = any(header in next_line.lower() for header in section_headers) or \
                                              next_line.isupper() and len(next_line) <= 50
                                next_is_job = any(char.isdigit() for char in next_line) and ('–' in next_line or '-' in next_line)
                                next_is_list = next_line.startswith(('●', '○', '§', '•', '-', '—', '|'))
                                
                                if next_is_header or next_is_job or next_is_list:
                                    break  # Start a new paragraph
                                
                                # Join this line to the current paragraph
                                paragraph += ' ' + next_line
                                i += 1
                            
                            processed_lines.append(paragraph)
                    
                    text_content = '\n'.join(processed_lines)
                    
                    # Fix specific known word splits that are common in PDF extraction
                    text_content = re.sub(r'Oper\s+ations', r'Operations', text_content)  # Fix "Oper ations"
                    text_content = re.sub(r'compr\s+ehensiv\s+e', r'comprehensive', text_content)  # Fix "compr ehensiv e"
                    text_content = re.sub(r'envir\s+onments', r'environments', text_content)  # Fix "envir onments"
                    text_content = re.sub(r'deliv\s+ery', r'delivery', text_content)  # Fix "deliv ery"
                    text_content = re.sub(r'procedur\s+es', r'procedures', text_content)  # Fix "procedur es"
                    text_content = re.sub(r'work\s+ﬂow', r'workflow', text_content)  # Fix "workﬂow" with special character
                    text_content = re.sub(r'conﬁgurations', r'configurations', text_content)  # Fix special character 'ﬁ'
                    text_content = re.sub(r'certiﬁed', r'certified', text_content)  # Fix special character 'ﬁ'
                    text_content = re.sub(r'certiﬁcations', r'certifications', text_content)  # Fix special character 'ﬁ'
                    
                    # Fix common concatenated words that appear in PDF extraction
                    text_content = re.sub(r'Manage([a-z]+)team', r'Manage \1 team', text_content)  # Fix "Managedateam"
                    text_content = re.sub(r'aregulated', r'a regulated', text_content)  # Fix "aregulated"
                    text_content = re.sub(r'acentralized', r'a centralized', text_content)  # Fix "acentralized"
                    
                    # Fix other common concatenated words
                    text_content = re.sub(r'deploy([a-z]+)om', r'deploy \1 om', text_content)  # Fix "deploycust om"
                    
                    # Fix common patterns where spaces were incorrectly inserted in the middle of words
                    text_content = re.sub(r'(\w{2,})-\s+(\w{2,})', r'\1-\2', text_content)  # Fix "high-v olume" -> "high-volume"
                    text_content = re.sub(r'(\w{2,})\s+([a-z])\s+(\w{2,})', r'\1\2\3', text_content)  # Fix "oper ations" -> "operations"
                    
                    # Fix specific issues from current output
                    text_content = re.sub(r'cust\s+om', r'custom', text_content)  # Fix "cust om" -> "custom"
                    text_content = re.sub(r'Manage\s+da', r'Managed a', text_content)  # Fix "Manage da" -> "Managed a"
                    
                    # Handle general case of single character splits within words
                    text_content = re.sub(r'(\w{3,})\s+(\w)\s+(\w{3,})', r'\1\2\3', text_content)  # Fix word-char-word pattern like "compr e hensive"
                    
                    # Fix any remaining two-word concatenations that should be separated
                    text_content = re.sub(r'([a-z])([A-Z])', r'\1 \2', text_content)  # Insert space between lowercase and uppercase: 'deliveryProven' -> 'delivery Proven'
                    
                    # Final cleanup: fix any remaining obvious word splits that we can identify
                    text_content = text_content.replace('aregulated', 'a regulated')  # Fix "aregulated"
                    text_content = text_content.replace('acentralized', 'a centralized')  # Fix "acentralized"
                    text_content = text_content.replace('deploycust om', 'deploy cust om')  # Fix "deploycust om"
                    text_content = text_content.replace('Saa S', 'SaaS')  # Fix "Saa S"
                    text_content = text_content.replace('Compu Com', 'CompuCom')  # Fix "Compu Com"
                    text_content = text_content.replace('Accu Code', 'AccuCode')  # Fix "Accu Code"
                    text_content = text_content.replace('Code Louisville', 'CodeLouisville')  # Fix "Code Louisville"
                    text_content = text_content.replace('Comp TIA', 'CompTIA')  # Fix "Comp TIA"
                    
                    # Additional fixes for remaining word splits
                    text_content = re.sub(r'infrastructur\s+es', r'infrastructure', text_content)
                    text_content = re.sub(r'high-v\s+olume', r'high-volume', text_content)
                    text_content = re.sub(r'oper\s+ations', r'operations', text_content)
                    text_content = re.sub(r'oper\s+ational', r'operational', text_content)
                    text_content = re.sub(r'high-le\s+verage', r'high-leverage', text_content)
                    text_content = re.sub(r'oppor\s+tunities', r'opportunities', text_content)
                    text_content = re.sub(r'deliv\s+ering', r'delivering', text_content)
                    text_content = re.sub(r'time-t\s+o-mill', r'time-to-mill', text_content)
                    text_content = re.sub(r'workﬂow', r'workflow', text_content)  # Fix special character
                    text_content = re.sub(r'Py\s+Installer', r'PyInstaller', text_content)
                    text_content = re.sub(r'ina\s+regulated', r'in a regulated', text_content)
                    text_content = re.sub(r'Createda', r'Created a', text_content)
                    
                    # Fix additional word splits seen in the current output
                    text_content = re.sub(r'exper\s+tise', r'expertise', text_content)  # Fix "exper tise"
                    text_content = re.sub(r'CAP\s+ABILI\s+TIES', r'CAPABILITIES', text_content)  # Fix "CAP ABILI TIES"
                    text_content = re.sub(r'PROJEC\s+TS', r'PROJECTS', text_content)  # Fix "PROJEC TS"
                    text_content = re.sub(r'certiﬁ\s+cations', r'certifications', text_content)  # Fix "certiﬁ cations" with special character
                    text_content = re.sub(r'Governance\s+Pro\s+jects', r'Governance Projects', text_content)  # Fix "Governance Projects" splits
                    text_content = re.sub(r'\s+\s+', ' ', text_content)  # Replace multiple spaces with single space
                    
                    # Replace multiple consecutive newlines with a single newline for remaining cases
                    text_content = re.sub(r'\n+', '\n', text_content)
                    # Fix any remaining excessive spacing
                    text_content = re.sub(r'\n\s+\n', '\n', text_content)
                    
                    # Save as text file
                    txt_path = OUTPUT_PATH / f"{name}.txt"
                    with open(txt_path, 'w', encoding='utf-8') as f:
                        f.write(text_content)
                    
                    file_path = str(txt_path)
                    
                    # ADDED: Warn about formatting loss
                    self._log_message(
                        f"PDF processed: {name} - extracted {len(text_content)} chars. "
                        f"Note: Formatting will be simplified in preview",
                        "info"
                    )
                except Exception as e:
                    self._log_message(f"PDF processing error: {e}", "error")
                    return
            
            # Add to database and set as active
            self.resume_model.add_resume(file_path, name, is_active=True)
            self._refresh_resume_list()
            
            # Load the preview with the processed content
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                    self.resume_preview.delete('1.0', tk.END)
                    
                    # Add content with improved visual formatting
                    lines = content.split('\n')
                    for i, line in enumerate(lines):
                        # Identify and highlight section headers
                        is_header = line.strip().upper() == line.strip() and len(line.strip()) < 50 and not line.strip().endswith('.') and len(line.strip()) > 0
                        
                        self.resume_preview.insert(tk.END, line + '\n')
                        if is_header:
                            # Tag headers to make them stand out
                            start_pos = self.resume_preview.index(f"{i+1}.0")
                            end_pos = self.resume_preview.index(f"{i+1}.end")
                            self.resume_preview.tag_add('header', start_pos, end_pos)
                    
                    # Configure header tags to make them stand out
                    self.resume_preview.tag_config('header', font=('Arial', 11, 'bold'), foreground='blue', spacing1=8, spacing3=8)
                    
                    # Add padding around the content
                    self.resume_preview.config(padx=20, pady=20)
            except Exception as e:
                self._log_message(f"Error loading resume preview after upload: {e}", "error")
            
            self._log_message(f"Resume uploaded successfully: {name} (set as active)", "info")
            
        except Exception as e:
            self._log_message(f"Error uploading resume: {e}", "error")
    
    def delete_selected_resume(self):
        """Delete the selected resume"""
        selection = self.resume_tree.selection()
        if not selection:
            messagebox.showwarning("Warning", "Please select a resume to delete")
            return
        
        # Confirm deletion
        if not messagebox.askyesno("Confirm Delete", "Are you sure you want to delete this resume?"):
            return
        
        try:
            item = self.resume_tree.item(selection[0])
            resume_path = item['values'][1]
            
            # Delete from database
            self.resume_model.delete_resume_by_path(resume_path)
            
            # Refresh list
            self._refresh_resume_list()
            self.resume_preview.delete('1.0', tk.END)
            
            self._log_message("Resume deleted successfully", "info")
            
        except Exception as e:
            self._log_message(f"Error deleting resume: {e}", "error")
    
    def set_active_resume(self):
        """Set the selected resume as active"""
        selection = self.resume_tree.selection()
        if not selection:
            messagebox.showwarning("Warning", "Please select a resume to set as active")
            return
        
        try:
            item = self.resume_tree.item(selection[0])
            resume_path = item['values'][1]
            
            # Set as active in database
            self.resume_model.set_active_resume_by_path(resume_path)
            
            # Refresh list
            self._refresh_resume_list()
            
            self._log_message(f"Active resume set to: {item['values'][0]}", "info")
            
        except Exception as e:
            self._log_message(f"Error setting active resume: {e}", "error")
    
    def _load_selected_resume(self):
        """Load the currently active resume"""
        try:
            active_resume = self.resume_model.get_active_resume()
            if not active_resume:
                self._log_message("No active resume found. Please upload a resume first.", "error")
                return None
            
            resume_path = active_resume['file_path']
            
            with open(resume_path, 'r', encoding='utf-8') as f:
                return f.read()
                
        except Exception as e:
            self._log_message(f"Error loading resume: {e}", "error")
            return None
    
    def _refresh_resume_list(self):
        """Refresh the resume list in treeview"""
        # Clear existing items
        for item in self.resume_tree.get_children():
            self.resume_tree.delete(item)
        
        # Load resumes from database
        resumes = self.resume_model.list_resumes()
        for resume in resumes:
            active_status = "Yes" if resume['is_active'] else "No"
            self.resume_tree.insert('', tk.END, values=(
                resume['name'],
                resume['file_path'],
                active_status
            ))
        
        # Set active resume path
        active_resume = self.resume_model.get_active_resume()
        if active_resume:
            self.active_resume_id = active_resume['id']
    
    def analyze_match(self):
        """Analyze resume-job compatibility and display match score"""
        self._log_message("Starting match analysis...", "info")
        
        # Get active resume
        resume_text = self._load_selected_resume()
        if not resume_text:
            messagebox.showerror("Error", "No active resume loaded. Please upload and set a resume as Active first.")
            self._log_message("Match analysis failed: No active resume", "error")
            return
        
        # Get job description (ONLY requirement for match analysis - Golden Rule #1)
        job_description = self.job_desc_text.get('1.0', tk.END).strip()
        if not job_description or len(job_description) < 100:
            messagebox.showerror("Error", "Please enter a detailed job description (minimum 100 characters).")
            self._log_message("Match analysis failed: Job description too short", "error")
            return
        
        # Golden Rule #1: Job title and company are NOT required for compatibility checking
        
        try:
            # Call AI match analyzer
        # Note: AI scores may vary slightly between runs due to the non-deterministic nature of language models
            # Show visual feedback that analysis is in progress
            original_text = self.analyze_button.cget("text")
            self.analyze_button.config(text="Analyzing...", state='disabled')
            self.status_label.config(text="Analyzing match with AI... Please wait")
            self.master.update_idletasks()
            
            self.match_data = analyze_match(resume_text, job_description)
            score = self.match_data.get('overall_score', 0)
            
            # Restore button and show results
            self.analyze_button.config(text=original_text, state='normal')
            self.status_label.config(text="Analysis complete")
            
            # Update match display
            self.match_label.config(text=f"Match Score: {score}%")
            
            # Color coding based on score
            if score >= self.current_threshold:
                self.match_label.config(foreground="green")
                self.start_button.config(state='normal')
                message = f"Strong match ({score}%)! Enter job title/company and click 'Start Tailoring' to proceed."
                self._log_message(f"Match analysis complete: {score}% (threshold met)", "info")
            else:
                self.match_label.config(foreground="red")
                self.start_button.config(state='disabled')
                message = f"Match score {score}% is below threshold ({self.current_threshold}%). Improve resume or consider different role."
                self._log_message(f"Match analysis complete: {score}% (below threshold)", "warning")
            
            # Show detailed breakdown
            self._show_match_details()
            
            # Show summary message
            messagebox.showinfo("Match Analysis", message)
            
            self.status_label.config(text="Ready")
            
        except Exception as e:
            self._log_message(f"Match analysis error: {e}", "error")
            messagebox.showerror("Error", f"Match analysis failed: {e}")
            self.status_label.config(text="Ready")
            self.match_label.config(text="Error during analysis", foreground="red")

    def _show_match_details(self):
        """Display detailed match breakdown in a popup window"""
        if not self.match_data:
            return
        
        details_window = tk.Toplevel(self.master)
        details_window.title("Match Analysis Details")
        details_window.geometry("700x500")
        
        # Create scrollable text area
        details_text = scrolledtext.ScrolledText(details_window, width=80, height=25, wrap=tk.WORD)
        details_text.pack(padx=10, pady=10, fill=tk.BOTH, expand=True)
        
        # Build details content
        score = self.match_data.get('overall_score', 0)
        
        # Handle complex data structures from AI response
        skills_data = self.match_data.get('skills_match', {})
        exp_data = self.match_data.get('experience_match', {})
        keywords_data = self.match_data.get('keywords_match', {})
        
        # Extract scores and analysis
        skills_score = skills_data.get('score', 'N/A') if isinstance(skills_data, dict) else skills_data
        exp_score = exp_data.get('score', 'N/A') if isinstance(exp_data, dict) else exp_data
        keywords_score = keywords_data.get('score', 'N/A') if isinstance(keywords_data, dict) else keywords_data
        
        # Extract analysis text
        skills_analysis = skills_data.get('analysis', []) if isinstance(skills_data, dict) else []
        exp_analysis = exp_data.get('analysis', []) if isinstance(exp_data, dict) else []
        keywords_analysis = keywords_data.get('analysis', []) if isinstance(keywords_data, dict) else []
        
        # Format analysis as readable text
        skills_text = '\n'.join([f'* {item}' if isinstance(item, str) else str(item) for item in skills_analysis]) if skills_analysis else 'No detailed analysis provided.'
        exp_text = '\n'.join([f'* {item}' if isinstance(item, str) else str(item) for item in exp_analysis]) if exp_analysis else 'No detailed analysis provided.'
        keywords_text = '\n'.join([f'* {item}' if isinstance(item, str) else str(item) for item in keywords_analysis]) if keywords_analysis else 'No detailed analysis provided.'
        
        # Get lists
        strengths = self.match_data.get('strengths', ['No strengths identified'])
        gaps = self.match_data.get('gaps', ['No gaps identified'])
        recommendations = self.match_data.get('recommendations', ['No recommendations'])
        
        # Format lists
        strengths_text = '\n'.join([f'* {item}' if isinstance(item, str) else str(item) for item in strengths])
        gaps_text = '\n'.join([f'* {item}' if isinstance(item, str) else str(item) for item in gaps])
        recommendations_text = '\n'.join([f'* {item}' if isinstance(item, str) else str(item) for item in recommendations])
        
        details = f"""MATCH SUMMARY
=============
Overall Score: {score}%
Skills Match: {skills_score}%
Experience Match: {exp_score}%
Keywords Match: {keywords_score}%

SKILLS ANALYSIS:
===============
{skills_text}

EXPERIENCE ANALYSIS:
==================
{exp_text}

KEYWORDS ANALYSIS:
=================
{keywords_text}

STRENGTHS:
==========
{strengths_text}

GAPS:
=====
{gaps_text}

RECOMMENDATIONS:
================
{recommendations_text}
"""
        
        details_text.insert('1.0', details)
        details_text.config(state='disabled')
        
        # Add close button
        close_button = ttk.Button(details_window, text="Close", command=details_window.destroy)
        close_button.pack(pady=5)
    
    def start_tailoring(self):
        """Validate all prerequisites and start AI-powered tailoring"""
        self._log_message("Validating tailoring prerequisites...", "info")
        
        # Get job description first
        job_description = self.job_desc_text.get('1.0', tk.END).strip()
        if not job_description or len(job_description) < 100:
            messagebox.showerror("Insufficient Job Description", "Please enter a detailed job description (minimum 100 characters).")
            return
        
        # Prerequisites: job title, company, description, match score >= threshold
        job_title = self.job_title_entry.get().strip()
        company = self.company_entry.get().strip()
        
        # If job title or company are missing, try to extract them using AI
        if not job_title or not company:
            self._log_message("Attempting to auto-extract job title and company...", "info")
            try:
                extracted_details = extract_job_details(job_description)
                if not job_title and extracted_details.get('job_title') and extracted_details['job_title'] != 'Unknown':
                    job_title = extracted_details['job_title']
                    self.job_title_entry.delete(0, tk.END)
                    self.job_title_entry.insert(0, job_title)
                    self._log_message(f"Auto-extracted job title: {job_title}", "info")
                
                if not company and extracted_details.get('company_name') and extracted_details['company_name'] != 'Unknown':
                    company = extracted_details['company_name']
                    self.company_entry.delete(0, tk.END)
                    self.company_entry.insert(0, company)
                    self._log_message(f"Auto-extracted company: {company}", "info")
            except Exception as e:
                self._log_message(f"Failed to auto-extract job details: {e}", "warning")
        
        # Validate that we now have both job title and company
        if not job_title:
            messagebox.showerror("Missing Job Title", "Please enter a job title for file naming.")
            return
        
        if not company:
            messagebox.showerror("Missing Company", "Please enter a company name for file naming.")
            return
        
        # Load resume
        resume_text = self._load_selected_resume()
        if not resume_text:
            messagebox.showerror("Missing Active Resume", "No active resume found. Please upload and set a resume as Active.")
            return
        
        # Check match data exists and meets threshold
        if not hasattr(self, 'match_data') or not self.match_data:
            messagebox.showerror("No Match Analysis", "Please click 'Analyze Match' first to check compatibility.")
            return
        
        score = self.match_data.get('overall_score', 0)
        if score < self.current_threshold:
            messagebox.showerror("Match Too Low", f"Match score {score}% is below minimum threshold of {self.current_threshold}%. Consider improving your resume or applying to a different role.")
            self._log_message(f"Tailoring blocked: match {score}% < threshold {self.current_threshold}%", "warning")
            return
        
        # All validations passed - proceed with AI tailoring
        self.set_ui_enabled(False)
        self._processing = True  # Track processing state
        self.status_label.config(text="Processing... Please wait", foreground="orange")
        self._log_message("Starting AI-powered tailoring...", "info")
        
        # Get role level
        role_level = self.role_var.get()
        
        # Start tailoring thread with AI engine
        thread = threading.Thread(
            target=self.tailor_application_thread,
            args=(job_title, company, job_description, self.job_url_entry.get(), resume_text, role_level, None)
        )
        thread.daemon = True
        thread.start()
    
    def tailor_application_thread(self, job_title, company, job_description, job_url, resume_text, role_level, custom_prompt):
        """Thread function for tailoring process"""
        try:
            # Process and tailor
            result = process_and_tailor_from_gui(
                resume_text, job_description, OUTPUT_PATH, 
                role_level, custom_prompt
            )
            
            # Check for required fields
            if not result or not result.get("resume_text") or not result.get("cover_letter"):
                raise Exception("AI returned incomplete or empty tailoring results")
            
            # FIXED: Use "info" instead of "success" for logging
            self._log_message("Resume tailoring completed successfully", "info")
            
            # Add to queue for main thread
            self.tailoring_queue.put({
                'status': 'success',
                'result': result,
                'job_title': job_title,
                'company': company,
                'job_description': job_description
            })
            
        except Exception as e:
            self._log_message(f"Error during tailoring: {e}", "error")
            self.tailoring_queue.put({
                'status': 'error',
                'error': str(e)
            })
    
    def on_tailoring_complete(self, result_data):
        """Handle completion of tailoring process"""
        # Reset processing flag
        self._processing = False
        
        if result_data['status'] == 'error':
            self.status_label.config(text="Tailoring failed", foreground="red")
            messagebox.showerror("Error", f"Tailoring failed: {result_data['error']}")
            self.set_ui_enabled(True)
            return
        
        try:
            result = result_data['result']
            job_title = result_data['job_title']
            company = result_data['company']
            job_description = result_data.get('job_description', '')
            
            # Get match score if available
            match_score = 0
            if hasattr(self, 'match_data') and self.match_data:
                match_score = self.match_data.get('overall_score', 0)
            
            # Save outputs
            self.save_outputs(
                result['resume_text'],
                result['cover_letter'],
                job_title,
                company,
                job_description,
                match_score,
                self.match_data  # Pass the full match summary data
            )
            
            # Clear fields
            self.clear_fields()
            
            self.status_label.config(text="Ready", foreground="green")
            messagebox.showinfo("Success", "Resume tailoring completed! Files saved to output folder.")
            self._log_message("Files saved successfully", "info")
            
        except Exception as e:
            self.status_label.config(text="Save failed", foreground="red")
            messagebox.showerror("Save Error", f"Error saving files: {e}")
            self._log_message(f"Save error: {e}", "error")
            return
        finally:
            self.set_ui_enabled(True)
    
    def _create_import_tab(self):
        """Create the Job Import tab for importing from LinkedIn, email, etc."""
        tab = ttk.Frame(self.notebook)
        self.notebook.add(tab, text="Import Job")
        
        # Title
        title_label = ttk.Label(tab, text="Import Job Description", style='Title.TLabel')
        title_label.grid(row=0, column=0, columnspan=3, pady=(0, 15), sticky=tk.W)
        
        # Import source selection
        ttk.Label(tab, text="Import Source:", style='Bold.TLabel').grid(row=1, column=0, sticky=tk.W, pady=5)
        self.import_source_var = tk.StringVar(value="Plain Text")
        source_combo = ttk.Combobox(tab, textvariable=self.import_source_var, 
                                  values=["Plain Text", "LinkedIn HTML", "Email Content"], 
                                  state='readonly', width=20, font=('Arial', 10))
        source_combo.grid(row=1, column=1, sticky=tk.W, pady=5)
        
        # Source-specific instructions
        self.import_instructions = ttk.Label(tab, 
                                          text="Select import source type and paste the job content below. See instructions for each source type.", 
                                          wraplength=600, 
                                          justify=tk.LEFT,
                                          font=('Arial', 10))
        self.import_instructions.grid(row=2, column=0, columnspan=3, sticky=tk.W, pady=5)
        
        # Detailed instructions label
        instructions_detail = ttk.Label(tab, 
                                      text="\nHow to use Import Job:\n" +
                                            "• Plain Text: Copy and paste job description from any source\n" +
                                            "• LinkedIn HTML: Copy HTML source of LinkedIn job posting (Ctrl+U in browser)\n" +
                                            "• Email Content: Copy entire email content including subject line\n\n" +
                                            "After importing, the job details will automatically populate in the Job Management tab.",
                                      wraplength=600, 
                                      justify=tk.LEFT,
                                      font=('Arial', 9),
                                      foreground='gray')
        instructions_detail.grid(row=3, column=0, columnspan=3, sticky=tk.W, pady=5)
        
        # Update instructions label row
        self.import_instructions.grid(row=2, column=0, columnspan=3, sticky=tk.W, pady=5)
        
        # Bind to update instructions when source changes
        source_combo.bind('<<ComboboxSelected>>', self._update_import_instructions)
        
        # Text area for job description
        ttk.Label(tab, text="Job Content:", style='Bold.TLabel').grid(row=3, column=0, sticky=tk.W, pady=5)
        self.import_text = scrolledtext.ScrolledText(tab, width=80, height=20, wrap=tk.WORD, font=('Arial', 10))
        self.import_text.grid(row=4, column=0, columnspan=3, sticky=(tk.W, tk.E, tk.N, tk.S), pady=5)
        
        # Buttons
        button_frame = ttk.Frame(tab)
        button_frame.grid(row=5, column=0, columnspan=3, pady=15)
        
        self.parse_button = ttk.Button(button_frame, text="Parse & Import", command=self._parse_and_import_job)
        self.parse_button.grid(row=0, column=0, padx=5)
        
        self.clear_import_button = ttk.Button(button_frame, text="Clear", command=self._clear_import_fields)
        self.clear_import_button.grid(row=0, column=1, padx=5)
        
        # Add quit button
        quit_button = ttk.Button(button_frame, text="Quit", command=self.quit_application)
        quit_button.grid(row=0, column=2, padx=5)
        
        # Configure grid weights
        tab.columnconfigure(0, weight=1)
        tab.rowconfigure(4, weight=1)
    
    def _update_import_instructions(self, event=None):
        """Update import instructions based on selected source"""
        source = self.import_source_var.get()
        
        instructions = {
            "Plain Text": "Paste job description text below:",
            "LinkedIn HTML": "Paste the HTML source of the LinkedIn job posting below:\n(Use Ctrl+U or View Source in your browser)",
            "Email Content": "Paste the raw email content below:\n(Include subject line and body)"
        }
        
        self.import_instructions.config(text=instructions.get(source, instructions["Plain Text"]))
    
    def _parse_and_import_job(self):
        """Parse imported job content and populate job fields"""
        try:
            # Import the job parser
            from job_parser import parse_linkedin_job_description, parse_email_job_description, parse_plain_text_job_description
            
            # Get content
            content = self.import_text.get('1.0', tk.END).strip()
            if not content:
                messagebox.showwarning("Import Error", "Please paste job content first.")
                return
            
            # Parse based on selected source
            source = self.import_source_var.get()
            
            if source == "LinkedIn HTML":
                job_data = parse_linkedin_job_description(content)
            elif source == "Email Content":
                job_data = parse_email_job_description(content)
            else:  # Plain Text
                job_data = parse_plain_text_job_description(content)
            
            # Switch to Job Management tab and populate fields
            self.notebook.select(0)  # First tab is Job Management
            
            # Populate fields
            self.job_title_entry.delete(0, tk.END)
            self.job_title_entry.insert(0, job_data.get('title', ''))
            
            self.company_entry.delete(0, tk.END)
            self.company_entry.insert(0, job_data.get('company', ''))
            
            self.job_desc_text.delete('1.0', tk.END)
            self.job_desc_text.insert('1.0', job_data.get('description', content))
            
            # Show success message
            messagebox.showinfo("Import Successful", 
                              f"Job imported successfully!\n\nTitle: {job_data.get('title', 'N/A')}\nCompany: {job_data.get('company', 'N/A')}")
            
            self._log_message(f"Job imported from {source}", "info")
            
        except Exception as e:
            messagebox.showerror("Import Error", f"Failed to parse job content: {str(e)}")
            self._log_message(f"Import error: {e}", "error")
    
    def _clear_import_fields(self):
        """Clear import fields"""
        self.import_text.delete('1.0', tk.END)
        self._log_message("Cleared import fields", "info")
    
    def _create_custom_prompt_tab(self):
        """Create the Custom Prompt Management tab."""
        tab = ttk.Frame(self.notebook)
        self.notebook.add(tab, text="Custom Prompts")
        
        # Title
        title_label = ttk.Label(tab, text="Custom Prompt Templates", style='Title.TLabel')
        title_label.grid(row=0, column=0, columnspan=4, pady=(0, 15), sticky=tk.W)
        
        # Instructions
        instruction_text = "Create custom prompt templates for specific roles or companies. Save them in prompts/user/ directory as .txt.j2 files."
        
        instruction_label = ttk.Label(tab, text=instruction_text, wraplength=600, justify=tk.LEFT, font=('Arial', 10))
        instruction_label.grid(row=1, column=0, columnspan=4, sticky=tk.W, pady=5)
        
        # Available variables section
        variables_frame = ttk.LabelFrame(tab, text="Available Template Variables", padding="10")
        variables_frame.grid(row=2, column=0, columnspan=4, sticky=(tk.W, tk.E), pady=5)
        
        variables_text = (
            "{{ role_level }} - The selected role level (Standard, Senior, Lead, Principal)\n"
            "{{ company_name }} - The company name from the job posting\n"
            "{{ job_title }} - The job title from the job posting\n"
            "{{ job_description }} - The full job description text\n"
            "{{ resume_text }} - The original resume text"
        )
        variables_label = ttk.Label(variables_frame, text=variables_text, font=('Arial', 9), justify=tk.LEFT)
        variables_label.grid(row=0, column=0, sticky=tk.W)
        
        # Template examples section
        examples_frame = ttk.LabelFrame(tab, text="Built-in Template Examples", padding="10")
        examples_frame.grid(row=3, column=0, columnspan=4, sticky=(tk.W, tk.E), pady=5)
        
        examples_text = (
            "Standard Template: prompts/system/system.txt.j2\n"
            "Senior Template: prompts/system/senior.txt.j2\n\n"
            "Load these examples to see how to structure your custom prompts."
        )
        examples_label = ttk.Label(examples_frame, text=examples_text, font=('Arial', 9), justify=tk.LEFT)
        examples_label.grid(row=0, column=0, sticky=tk.W)
        
        # Prompt name entry
        ttk.Label(tab, text="Prompt Name:", style='Bold.TLabel').grid(row=4, column=0, sticky=tk.W, pady=5)
        self.prompt_name_entry = ttk.Entry(tab, width=40, font=('Arial', 10))
        self.prompt_name_entry.grid(row=4, column=1, columnspan=2, sticky=(tk.W, tk.E), pady=5)
        self.prompt_name_entry.insert(0, "my_custom_prompt")
        
        # Template editor
        ttk.Label(tab, text="Prompt Template:", style='Bold.TLabel').grid(row=5, column=0, sticky=tk.W, pady=5)
        self.prompt_text = scrolledtext.ScrolledText(tab, width=80, height=15, wrap=tk.WORD, font=('Arial', 10))
        self.prompt_text.grid(row=5, column=1, columnspan=2, sticky=(tk.W, tk.E, tk.N, tk.S), pady=5)
        
        # Load default template
        try:
            with open("prompts/user/custom_template.txt.j2", "r") as f:
                self.prompt_text.insert("1.0", f.read())
        except:
            default_template = """# Custom Resume Tailoring Prompt Template
# Available variables: {{ role_level }}, {{ company_name }}, {{ job_title }}, {{ job_description }}, {{ resume_text }}

You are an expert resume tailor specializing in {{ role_level }} positions for {{ company_name }}.

TASK: Tailor the following resume for the {{ job_title }} role.

Original Resume:
{{ resume_text }}

Job Description:
{{ job_title }} at {{ company_name }}
{{ job_description }}

Requirements:
1. Match language to job description keywords
2. Quantify achievements where possible
3. Emphasize relevant certifications and experience
4. Optimize for ATS (Applicant Tracking Systems)
5. Maintain professional tone appropriate for {{ role_level }} level

Output ONLY the tailored resume followed by a cover letter.

Format your response exactly as follows:
[TAILORING_COMPLETE]
[Tailored resume content here]
[COVER LETTER]
[Cover letter content here]
[END_APPLICATION_MATERIALS]
"""
            self.prompt_text.insert("1.0", default_template)
        
        # Buttons
        button_frame = ttk.Frame(tab)
        button_frame.grid(row=6, column=0, columnspan=4, pady=15)
        
        save_button = ttk.Button(button_frame, text="Save Prompt", command=self.save_custom_prompt)
        save_button.grid(row=0, column=0, padx=5)
        
        load_button = ttk.Button(button_frame, text="Load Prompt", command=self.load_custom_prompt)
        load_button.grid(row=0, column=1, padx=5)
        
        load_example_button = ttk.Button(button_frame, text="Load Example", command=self.load_example_prompt)
        load_example_button.grid(row=0, column=2, padx=5)
        
        preview_button = ttk.Button(button_frame, text="Preview Variables", command=self.preview_variables)
        preview_button.grid(row=0, column=3, padx=5)
        
        clear_button = ttk.Button(button_frame, text="Clear", command=self.clear_prompt_editor)
        clear_button.grid(row=0, column=4, padx=5)
        
        # Add quit button
        quit_button = ttk.Button(button_frame, text="Quit", command=self.quit_application)
        quit_button.grid(row=0, column=5, padx=5)
        
        # Make text area expandable
        tab.columnconfigure(1, weight=1)
        tab.rowconfigure(5, weight=1)
    
    def save_custom_prompt(self):
        """Save custom prompt template to file."""
        prompt_name = self.prompt_name_entry.get().strip()
        if not prompt_name.endswith('.txt.j2'):
            prompt_name += '.txt.j2'
        
        if not prompt_name or prompt_name == '.txt.j2':
            messagebox.showerror("Error", "Please enter a valid prompt name.")
            return
        
        template_content = self.prompt_text.get('1.0', tk.END).strip()
        
        try:
            with open(f"prompts/user/{prompt_name}", 'w') as f:
                f.write(template_content)
            messagebox.showinfo("Success", f"Prompt saved as {prompt_name}")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to save prompt: {e}")
    
    def load_custom_prompt(self):
        """Load custom prompt template using curated dropdown (Golden Rule #5 - Security)"""
        prompts_dir = Path("prompts/user")
        
        # Verify directory exists
        if not prompts_dir.exists():
            prompts_dir.mkdir(parents=True, exist_ok=True)
            messagebox.showwarning("No Prompts", "No custom prompts found. Created prompts/user/ directory. Please add .txt.j2 files.")
            self._log_message("Prompts directory created but empty", "warning")
            return
        
        # Get available prompts (no filesystem navigation - Golden Rule #5)
        try:
            available_prompts = [f.name for f in prompts_dir.glob("*.txt.j2")]
        except Exception as e:
            self._log_message(f"Error reading prompts directory: {e}", "error")
            messagebox.showerror("Error", f"Could not read prompts directory: {e}")
            return
        
        if not available_prompts:
            messagebox.showwarning("No Prompts", "No custom prompts found in prompts/user/. Please add .txt.j2 files.")
            self._log_message("No prompts available in prompts/user/", "warning")
            return
        
        # Create selection window (curated list - no file dialog)
        selection_window = tk.Toplevel(self.master)
        selection_window.title("Select Custom Prompt")
        selection_window.geometry("400x300")
        
        ttk.Label(selection_window, text="Available Prompts:", font=('Arial', 10, 'bold')).pack(pady=10)
        
        # Create listbox with available prompts
        listbox = tk.Listbox(selection_window, selectmode=tk.SINGLE, height=10, width=40)
        for prompt in sorted(available_prompts):
            listbox.insert(tk.END, prompt)
        listbox.pack(pady=10, padx=10, fill=tk.BOTH, expand=True)
        
        # Add selection button
        def select_prompt():
            selection = listbox.curselection()
            if not selection:
                messagebox.showwarning("No Selection", "Please select a prompt.")
                return
            
            prompt_name = listbox.get(selection[0])
            selection_window.destroy()
            
            # Load the selected prompt
            try:
                prompt_path = prompts_dir / prompt_name
                with open(prompt_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                
                self.prompt_name_entry.delete(0, tk.END)
                self.prompt_name_entry.insert(0, prompt_name.replace('.txt.j2', ''))
                
                self.prompt_text.delete('1.0', tk.END)
                self.prompt_text.insert("1.0", content)
                
                self._log_message(f"Loaded custom prompt: {prompt_name}", "info")
                messagebox.showinfo("Success", f"Loaded prompt: {prompt_name}")
                
            except Exception as e:
                self._log_message(f"Load prompt error: {e}", "error")
                messagebox.showerror("Error", f"Failed to load prompt: {e}")
        
        ttk.Button(selection_window, text="Load Selected", command=select_prompt).pack(pady=10)
    
    def clear_prompt_editor(self):
        """Clear the prompt editor fields"""
        self.prompt_name_entry.delete(0, tk.END)
        self.prompt_text.delete('1.0', tk.END)
        self._log_message("Cleared prompt editor", "info")
    
    def load_example_prompt(self):
        """Load example prompt templates from system directory"""
        examples_dir = Path("prompts/system")
        
        # Verify directory exists
        if not examples_dir.exists():
            messagebox.showwarning("No Examples", "System prompt examples not found.")
            return
        
        # Get available examples
        try:
            available_examples = [f.name for f in examples_dir.glob("*.txt.j2")]
        except Exception as e:
            self._log_message(f"Error reading examples directory: {e}", "error")
            messagebox.showerror("Error", f"Could not read examples directory: {e}")
            return
        
        if not available_examples:
            messagebox.showwarning("No Examples", "No system prompt examples found.")
            return
        
        # Create selection window
        selection_window = tk.Toplevel(self.master)
        selection_window.title("Select Example Prompt")
        selection_window.geometry("400x300")
        
        ttk.Label(selection_window, text="Available Examples:", style='Bold.TLabel').pack(pady=10)
        
        # Create listbox with available examples
        listbox = tk.Listbox(selection_window, selectmode=tk.SINGLE, height=10, width=40)
        for example in sorted(available_examples):
            listbox.insert(tk.END, example)
        listbox.pack(pady=10, padx=10, fill=tk.BOTH, expand=True)
        
        # Add selection button
        def select_example():
            selection = listbox.curselection()
            if not selection:
                messagebox.showwarning("No Selection", "Please select an example.")
                return
            
            example_name = listbox.get(selection[0])
            selection_window.destroy()
            
            # Load the selected example
            try:
                example_path = examples_dir / example_name
                with open(example_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                
                self.prompt_name_entry.delete(0, tk.END)
                self.prompt_name_entry.insert(0, example_name.replace('.txt.j2', ''))
                
                self.prompt_text.delete('1.0', tk.END)
                self.prompt_text.insert("1.0", content)
                
                self._log_message(f"Loaded example prompt: {example_name}", "info")
                messagebox.showinfo("Success", f"Loaded example: {example_name}")
                
            except Exception as e:
                self._log_message(f"Load example error: {e}", "error")
                messagebox.showerror("Error", f"Failed to load example: {e}")
        
        ttk.Button(selection_window, text="Load Selected", command=select_example).pack(pady=10)
    
    def preview_variables(self):
        """Preview how variables would be substituted in the current prompt"""
        # Get current prompt content
        prompt_content = self.prompt_text.get('1.0', tk.END)
        
        # Create preview window
        preview_window = tk.Toplevel(self.master)
        preview_window.title("Variable Preview")
        preview_window.geometry("600x400")
        
        # Sample values for demonstration
        sample_values = {
            "role_level": "Senior",
            "company_name": "TechCorp Inc.",
            "job_title": "Senior Software Engineer",
            "job_description": "We are looking for an experienced software engineer with 5+ years in Python and cloud technologies...",
            "resume_text": "John Doe\nSenior Software Engineer\nExperienced in Python, AWS, and distributed systems..."
        }
        
        # Show sample values
        ttk.Label(preview_window, text="Sample Variable Values:", style='Bold.TLabel').pack(pady=5)
        
        values_text = scrolledtext.ScrolledText(preview_window, width=70, height=8, wrap=tk.WORD)
        values_display = "\n".join([f"{{{{ {key} }}}} = {value}" for key, value in sample_values.items()])
        values_text.insert("1.0", values_display)
        values_text.pack(pady=5, padx=10, fill=tk.BOTH, expand=True)
        values_text.config(state=tk.DISABLED)
        
        # Show how prompt would look with substitutions
        ttk.Label(preview_window, text="Prompt with Sample Substitutions:", style='Bold.TLabel').pack(pady=5)
        
        preview_text = scrolledtext.ScrolledText(preview_window, width=70, height=10, wrap=tk.WORD)
        substituted_content = prompt_content
        for key, value in sample_values.items():
            substituted_content = substituted_content.replace(f"{{{{ {key} }}}}", f"[{value}]")
        preview_text.insert("1.0", substituted_content)
        preview_text.pack(pady=5, padx=10, fill=tk.BOTH, expand=True)
        preview_text.config(state=tk.DISABLED)
        
        ttk.Label(preview_window, text="Note: Brackets [] indicate where your actual values will be inserted", font=('Arial', 9)).pack(pady=5)
    
    def validate_prompt_template(self, template_content):
        """Validate that prompt template contains required structure and variables"""
        required_sections = [
            "[TAILORING_COMPLETE]",
            "[COVER LETTER]",
            "[END_APPLICATION_MATERIALS]"
        ]
        
        missing_sections = []
        for section in required_sections:
            if section not in template_content:
                missing_sections.append(section)
        
        if missing_sections:
            return False, f"Missing required sections: {', '.join(missing_sections)}"
        
        return True, "Template is valid"
    
    def save_custom_prompt(self):
        """Save custom prompt template to file."""
        prompt_name = self.prompt_name_entry.get().strip()
        if not prompt_name.endswith('.txt.j2'):
            prompt_name += '.txt.j2'
        
        if not prompt_name or prompt_name == '.txt.j2':
            messagebox.showerror("Error", "Please enter a valid prompt name.")
            return
        
        template_content = self.prompt_text.get('1.0', tk.END).strip()
        
        # Validate template before saving
        is_valid, validation_message = self.validate_prompt_template(template_content)
        if not is_valid:
            result = messagebox.askyesno(
                "Template Validation Warning", 
                f"{validation_message}\n\nDo you want to save anyway?"
            )
            if not result:
                return
        
        try:
            with open(f"prompts/user/{prompt_name}", 'w') as f:
                f.write(template_content)
            messagebox.showinfo("Success", f"Prompt saved as {prompt_name}")
            self._log_message(f"Custom prompt saved: {prompt_name}", "info")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to save prompt: {e}")
            self._log_message(f"Error saving custom prompt: {e}", "error")
    
    def _create_tailored_docs_tab(self):
        """Create the Tailored Documents tab with split view for job description, resume and cover letter"""
        tab = ttk.Frame(self.notebook)
        self.notebook.add(tab, text="Tailored Documents")
        
        # Title
        title_label = ttk.Label(tab, text="View & Export Tailored Applications", style='Title.TLabel')
        title_label.grid(row=0, column=0, columnspan=3, pady=(0, 15), sticky=tk.W)
        
        # Applications list
        ttk.Label(tab, text="Select Application:", style='Bold.TLabel').grid(row=1, column=0, sticky=tk.W, pady=5)
        
        # Create treeview for applications
        columns = ('Job Title', 'Company', 'Date')
        self.applications_tree = ttk.Treeview(tab, columns=columns, show='headings', height=8)
        self.applications_tree.heading('Job Title', text='Job Title')
        self.applications_tree.heading('Company', text='Company')
        self.applications_tree.heading('Date', text='Date')
        
        # Set column widths to ensure proper alignment and prevent visual interference
        self.applications_tree.column('Job Title', width=250, minwidth=150, anchor='w')
        self.applications_tree.column('Company', width=200, minwidth=100, anchor='w')
        self.applications_tree.column('Date', width=150, minwidth=100, anchor='w')
        
        # Configure treeview style to ensure clean column separation without interfering lines
        style = ttk.Style()
        style.configure('Treeview.Heading', font=('Arial', 10, 'bold'), relief='flat')
        # Ensure clean appearance without excessive separators
        style.configure('Treeview', rowheight=25, borderwidth=0, relief='flat')
        
        self.applications_tree.grid(row=2, column=0, columnspan=2, sticky=(tk.W, tk.E, tk.N, tk.S), pady=5)
        
        # Add scrollbar
        scrollbar = ttk.Scrollbar(tab, orient=tk.VERTICAL, command=self.applications_tree.yview)
        self.applications_tree.configure(yscrollcommand=scrollbar.set)
        scrollbar.grid(row=2, column=2, sticky=(tk.N, tk.S))
        
        # Bind selection event
        self.applications_tree.bind('<<TreeviewSelect>>', self._on_application_select)
        
        # Split view for job description, resume and cover letter
        ttk.Label(tab, text="Job Description", style='Bold.TLabel').grid(row=3, column=0, sticky=tk.W, pady=(15, 5))
        ttk.Label(tab, text="Tailored Resume", style='Bold.TLabel').grid(row=3, column=1, sticky=tk.W, pady=(15, 5))
        ttk.Label(tab, text="Cover Letter", style='Bold.TLabel').grid(row=3, column=2, sticky=tk.W, pady=(15, 5))
        
        # Text areas for job description, resume and cover letter
        self.job_description_text = scrolledtext.ScrolledText(tab, width=35, height=20, wrap=tk.WORD, font=('Arial', 10))
        self.job_description_text.grid(row=4, column=0, sticky=(tk.W, tk.E, tk.N, tk.S), padx=(0, 5), pady=5)
        
        self.tailored_resume_text = scrolledtext.ScrolledText(tab, width=35, height=20, wrap=tk.WORD, font=('Courier New', 10), padx=5, pady=5)
        self.tailored_resume_text.grid(row=4, column=1, sticky=(tk.W, tk.E, tk.N, tk.S), padx=(5, 5), pady=5)
        
        self.cover_letter_text = scrolledtext.ScrolledText(tab, width=35, height=20, wrap=tk.WORD, font=('Courier New', 10), padx=5, pady=5)
        self.cover_letter_text.grid(row=4, column=2, sticky=(tk.W, tk.E, tk.N, tk.S), padx=(5, 0), pady=5)
        
        # Export section
        export_frame = ttk.LabelFrame(tab, text="Export Options", padding="10")
        export_frame.grid(row=5, column=0, columnspan=3, pady=(15, 10), sticky=(tk.W, tk.E))
        export_frame.columnconfigure(1, weight=1)
        
        ttk.Label(export_frame, text="Format:", font=('Arial', 10)).grid(row=0, column=0, padx=(0, 5))
        self.export_format_var = tk.StringVar(value="PDF")
        export_format_combo = ttk.Combobox(export_frame, textvariable=self.export_format_var, 
                                         values=["PDF", "Word (.docx)", "Plain Text (.txt)", "ATS-Optimized"], 
                                         state='readonly', width=15, font=('Arial', 10))
        export_format_combo.grid(row=0, column=1, padx=5, sticky=tk.W)
        
        # Buttons
        button_frame = ttk.Frame(tab)
        button_frame.grid(row=6, column=0, columnspan=3, pady=10)
        
        self.refresh_apps_button = ttk.Button(button_frame, text="Refresh Applications", command=self._refresh_applications_list)
        self.refresh_apps_button.grid(row=0, column=0, padx=5)
        
        self.export_button = ttk.Button(button_frame, text="Export Documents", command=self._export_documents, state='disabled')
        self.export_button.grid(row=0, column=1, padx=5)
        
        self.delete_button = ttk.Button(button_frame, text="Delete Selected", command=self._delete_selected_application, state='disabled')
        self.delete_button.grid(row=0, column=2, padx=5)
        
        # Add quit button
        quit_button = ttk.Button(button_frame, text="Quit", command=self.quit_application)
        quit_button.grid(row=0, column=3, padx=5)
        
        # Refresh applications list
        self._refresh_applications_list()
        
        # Configure grid weights
        tab.columnconfigure(0, weight=1)
        tab.columnconfigure(1, weight=1)
        tab.columnconfigure(2, weight=1)
        tab.rowconfigure(2, weight=1)
        tab.rowconfigure(4, weight=2)
    
    def _refresh_applications_list(self):
        """Refresh the applications list in treeview"""
        # Clear existing items
        for item in self.applications_tree.get_children():
            self.applications_tree.delete(item)
        
        # Load applications from database
        applications = self.db_manager.get_all_applications()
        for app in applications:
            # Format date (handle both formats with and without microseconds)
            try:
                # Try with microseconds first
                dt = datetime.strptime(app['created_at'], '%Y-%m-%d %H:%M:%S.%f')
            except ValueError:
                # Fall back to without microseconds
                dt = datetime.strptime(app['created_at'], '%Y-%m-%d %H:%M:%S')
            # Format consistently with more readable format
            created_at = dt.strftime('%m/%d/%Y %H:%M')
            self.applications_tree.insert('', tk.END, values=(
                app['job_title'],
                app['company_name'],
                created_at
            ), iid=app['id'])
    
    def _on_application_select(self, event):
        """Handle application selection in treeview"""
        selection = self.applications_tree.selection()
        if selection:
            app_id = selection[0]
            # Get application details
            applications = self.db_manager.get_all_applications()
            selected_app = None
            for app in applications:
                if str(app["id"]) == app_id:
                    selected_app = app
                    break
            
            if selected_app:
                # Store match score for export functionality
                self.current_selected_app = selected_app
                
                # Load job description content
                try:
                    job_desc_path = selected_app.get("job_description_path")
                    # Handle both None and string 'None' cases
                    if job_desc_path and job_desc_path != 'None' and os.path.exists(job_desc_path):
                        with open(job_desc_path, "r", encoding="utf-8") as f:
                            job_description_content = f.read()
                        self.job_description_text.delete("1.0", tk.END)
                        self.job_description_text.insert("1.0", job_description_content)
                    else:
                        self.job_description_text.delete("1.0", tk.END)
                        self.job_description_text.insert("1.0", "Job description not available.")
                except Exception as e:
                    self.job_description_text.delete("1.0", tk.END)
                    self.job_description_text.insert("1.0", f"Error loading job description: {e}")
                
                # Load resume content
                try:
                    with open(selected_app["resume_path"], "r", encoding="utf-8") as f:
                        resume_content = f.read()
                    self.tailored_resume_text.delete("1.0", tk.END)
                    self.tailored_resume_text.insert("1.0", resume_content)
                except Exception as e:
                    self.tailored_resume_text.delete("1.0", tk.END)
                    self.tailored_resume_text.insert("1.0", f"Error loading resume: {e}")
                
                # Load cover letter content
                try:
                    with open(selected_app["cover_letter_path"], "r", encoding="utf-8") as f:
                        cover_letter_content = f.read()
                    self.cover_letter_text.delete("1.0", tk.END)
                    self.cover_letter_text.insert("1.0", cover_letter_content)
                except Exception as e:
                    self.cover_letter_text.delete("1.0", tk.END)
                    self.cover_letter_text.insert("1.0", f"Error loading cover letter: {e}")
                
                # Enable export and delete buttons
                self.export_button.config(state="normal")
                self.delete_button.config(state="normal")
            else:
                self.job_description_text.delete("1.0", tk.END)
                self.tailored_resume_text.delete("1.0", tk.END)
                self.cover_letter_text.delete("1.0", tk.END)
                self.export_button.config(state="disabled")
                self.delete_button.config(state="disabled")
        else:
            # Only clear if there's no selection, not during refresh
            # Disable export and delete buttons when no selection
            self.export_button.config(state="disabled")
            self.delete_button.config(state="disabled")

    
    def _export_documents(self):
        """Export current tailored documents in selected format"""
        if not hasattr(self, 'current_selected_app'):
            messagebox.showwarning("Export Error", "Please select an application first.")
            return
        
        # Get selected export format
        export_format = self.export_format_var.get()
        
        try:
            # Get job title and company for filename
            job_title = self.current_selected_app['job_title'].replace(' ', '_').replace('/', '_')
            company = self.current_selected_app['company_name'].replace(' ', '_').replace('/', '_')
            
            # Handle different export formats
            if export_format == "PDF":
                self._export_as_pdf(job_title, company)
            elif export_format == "Word (.docx)":
                self._export_as_word(job_title, company)
            elif export_format == "Plain Text (.txt)":
                self._export_as_text(job_title, company)
            elif export_format == "ATS-Optimized":
                self._export_as_ats_optimized(job_title, company)
            else:
                messagebox.showerror("Export Error", f"Unsupported export format: {export_format}")
        
        except Exception as e:
            messagebox.showerror("Export Error", f"Failed to export documents: {str(e)}")
            self._log_message(f"Export error: {e}", "error")
    
    def _export_as_pdf(self, job_title, company):
        """Export current tailored documents as PDF"""
        try:
            # Import reportlab here to avoid issues if not installed
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.enums import TA_CENTER, TA_LEFT
            
            # Ask user for save location
            file_path = filedialog.asksaveasfilename(
                defaultextension=".pdf",
                filetypes=[("PDF files", "*.pdf"), ("All files", "*.*")],
                initialfile=f"{company}_{job_title}_Application.pdf"
            )
            
            if not file_path:
                return
            
            # Create PDF document
            doc = SimpleDocTemplate(file_path, pagesize=letter)
            styles = getSampleStyleSheet()
            
            # Custom styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                alignment=TA_CENTER,
                fontSize=16,
                spaceAfter=30
            )
            
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=14,
                spaceBefore=20,
                spaceAfter=10
            )
            
            # Custom normal style with smaller bullet points
            normal_style = ParagraphStyle(
                'CustomNormal',
                parent=styles['Normal'],
                bulletFontSize=8,  # Smaller bullet points
                bulletIndent=10,
                leftIndent=20
            )
            
            # Build document content
            story = []
            
            # Tailored Resume
            story.append(Paragraph("Tailored Resume", heading_style))
            
            # Read resume content
            with open(self.current_selected_app['resume_path'], 'r', encoding='utf-8') as f:
                resume_content = f.read()
            
            # Add resume content
            resume_para = Paragraph(resume_content.replace('\n', '<br/>'), normal_style)
            story.append(resume_para)
            story.append(PageBreak())
            
            # Cover Letter
            story.append(Paragraph("Cover Letter", heading_style))
            
            # Read cover letter content
            with open(self.current_selected_app['cover_letter_path'], 'r', encoding='utf-8') as f:
                cover_letter_content = f.read()
            
            # Add cover letter content
            cover_letter_para = Paragraph(cover_letter_content.replace('\n', '<br/>'), normal_style)
            story.append(cover_letter_para)
            
            # Build PDF
            doc.build(story)
            
            messagebox.showinfo("PDF Export", f"Documents exported successfully to:\n{file_path}")
            self._log_message(f"PDF exported to: {file_path}", "info")
            
        except ImportError:
            messagebox.showerror("PDF Export Error", "ReportLab library not found. Please install it with: pip install reportlab")
            self._log_message("ReportLab not installed for PDF export", "error")
        except Exception as e:
            raise  # Re-raise to be caught by parent method
    
    def _export_as_word(self, job_title, company):
        """Export current tailored documents as Word document"""
        try:
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            # Ask user for save location
            file_path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word files", "*.docx"), ("All files", "*.*")],
                initialfile=f"{company}_{job_title}_Application.docx"
            )
            
            if not file_path:
                return
            
            # Create Word document
            doc = Document()
            
            # Add title
            title = doc.add_heading('Job Application Documents', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add company and job title info
            info_para = doc.add_paragraph()
            info_run = info_para.add_run(f'{company} - {job_title}')
            info_run.bold = True
            
            # Add tailored resume
            doc.add_heading('Tailored Resume', level=1)
            
            # Read resume content
            with open(self.current_selected_app['resume_path'], 'r', encoding='utf-8') as f:
                resume_content = f.read()
            
            # Add resume content as paragraphs
            for line in resume_content.split('\n'):
                if line.strip():
                    doc.add_paragraph(line)
            
            # Add page break
            doc.add_page_break()
            
            # Add cover letter
            doc.add_heading('Cover Letter', level=1)
            
            # Read cover letter content
            with open(self.current_selected_app['cover_letter_path'], 'r', encoding='utf-8') as f:
                cover_letter_content = f.read()
            
            # Add cover letter content as paragraphs
            for line in cover_letter_content.split('\n'):
                if line.strip():
                    doc.add_paragraph(line)
            
            # Save document
            doc.save(file_path)
            
            messagebox.showinfo("Word Export", f"Documents exported successfully to:\n{file_path}")
            self._log_message(f"Word document exported to: {file_path}", "info")
            
        except ImportError:
            messagebox.showerror("Word Export Error", "python-docx library not found. Please install it with: pip install python-docx")
            self._log_message("python-docx not installed for Word export", "error")
        except Exception as e:
            raise  # Re-raise to be caught by parent method
    
    def _export_as_text(self, job_title, company):
        """Export current tailored documents as plain text files"""
        try:
            # Ask user for save location for zip file
            file_path = filedialog.asksaveasfilename(
                defaultextension=".zip",
                filetypes=[("ZIP files", "*.zip"), ("All files", "*.*")],
                initialfile=f"{company}_{job_title}_Application.zip"
            )
            
            if not file_path:
                return
            
            import zipfile
            from io import StringIO
            
            # Create zip file with all documents
            with zipfile.ZipFile(file_path, 'w') as zipf:
                # Add resume
                with open(self.current_selected_app['resume_path'], 'r', encoding='utf-8') as f:
                    zipf.writestr(f"{job_title}_resume.txt", f.read())
                
                # Add cover letter
                with open(self.current_selected_app['cover_letter_path'], 'r', encoding='utf-8') as f:
                    zipf.writestr(f"{job_title}_cover_letter.txt", f.read())
                
                # Add job description if available
                job_desc_path = self.current_selected_app.get('job_description_path')
                if job_desc_path and job_desc_path != 'None' and os.path.exists(job_desc_path):
                    with open(job_desc_path, 'r', encoding='utf-8') as f:
                        zipf.writestr(f"{job_title}_job_description.txt", f.read())
            
            messagebox.showinfo("Text Export", f"Documents exported successfully to:\n{file_path}")
            self._log_message(f"Text files exported to ZIP: {file_path}", "info")
            
        except Exception as e:
            raise  # Re-raise to be caught by parent method
    
    def _export_as_ats_optimized(self, job_title, company):
        """Export current tailored documents as ATS-optimized plain text"""
        try:
            # Ask user for save location
            file_path = filedialog.asksaveasfilename(
                defaultextension=".txt",
                filetypes=[("Text files", "*.txt"), ("All files", "*.*")],
                initialfile=f"{company}_{job_title}_ATS_Optimized.txt"
            )
            
            if not file_path:
                return
            
            # Read all content
            with open(self.current_selected_app['resume_path'], 'r', encoding='utf-8') as f:
                resume_content = f.read()
            
            with open(self.current_selected_app['cover_letter_path'], 'r', encoding='utf-8') as f:
                cover_letter_content = f.read()
            
            # Create ATS-optimized content (remove special formatting, extra spaces, etc.)
            ats_content = f"""JOB APPLICATION FOR: {company} - {job_title}

TAILORING DATE: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

=====================
TAILORING SCORE: {self.current_selected_app.get('match_score', 'N/A')}
=====================

TAILORING NOTES:
This document has been optimized for Applicant Tracking Systems (ATS).
All formatting has been simplified for maximum compatibility.

=====================
RESUME
=====================
{resume_content}

=====================
COVER LETTER
=====================
{cover_letter_content}"""
            
            # Simplify for ATS (remove extra whitespace, normalize line endings)
            lines = [line.strip() for line in ats_content.split('\n')]
            ats_content = '\n'.join(lines)
            
            # Save to file
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(ats_content)
            
            messagebox.showinfo("ATS Export", f"ATS-optimized document exported successfully to:\n{file_path}")
            self._log_message(f"ATS-optimized text exported to: {file_path}", "info")
            
        except Exception as e:
            raise  # Re-raise to be caught by parent method
    
    def _delete_selected_application(self):
        """Delete the currently selected application and its associated files"""
        if not hasattr(self, 'current_selected_app'):
            messagebox.showwarning("Delete Error", "Please select an application first.")
            return
        
        # Confirm deletion type
        delete_choice = messagebox.askyesno(
            "Delete Type", 
            "Do you want to delete ALL files (including job description)?\n\n"
            "YES = Delete everything (cannot be undone)\n"
            "NO = Delete only resume and cover letter (job description preserved)"
            , icon='question', type='yesno')
        
        # If user cancels the dialog, return
        if delete_choice is None:
            return
        
        # delete_choice is True for 'yes' (delete all), False for 'no' (delete tailored docs only)
        delete_all = delete_choice
        
        try:
            app_id = self.current_selected_app['id']
            resume_path = self.current_selected_app['resume_path']
            cover_letter_path = self.current_selected_app['cover_letter_path']
            job_desc_path = self.current_selected_app.get('job_description_path')
            
            # Delete files based on user choice
            if delete_all:
                # Delete all files including job description
                files_to_delete = [resume_path, cover_letter_path]
                if job_desc_path and job_desc_path != 'None' and os.path.exists(job_desc_path):
                    files_to_delete.append(job_desc_path)
                delete_message = "All files including job description have been deleted."
            else:
                # Delete only tailored resume and cover letter files (preserve job description)
                files_to_delete = [resume_path, cover_letter_path]
                delete_message = "Tailored resume and cover letter have been deleted.\n\nThe job description has been preserved for future reference."
            
            for file_path in files_to_delete:
                try:
                    if os.path.exists(file_path):
                        os.remove(file_path)
                        self._log_message(f"Deleted file: {file_path}", "info")
                except Exception as e:
                    self._log_message(f"Error deleting file {file_path}: {e}", "error")
            
            # Delete from database
            self.db_manager.delete_application(app_id)
            
            # Clear the current selection and refresh the list
            self.current_selected_app = None
            self.job_description_text.delete("1.0", tk.END)
            self.tailored_resume_text.delete("1.0", tk.END)
            self.cover_letter_text.delete("1.0", tk.END)
            self.export_button.config(state="disabled")
            self.delete_button.config(state="disabled")
            self._refresh_applications_list()
            
            messagebox.showinfo("Delete Successful", delete_message)
            if delete_all:
                self._log_message("All files including job description deleted successfully", "info")
            else:
                self._log_message("Tailored resume and cover letter deleted successfully", "info")
            
        except Exception as e:
            messagebox.showerror("Delete Error", f"Failed to delete application: {str(e)}")
            self._log_message(f"Delete error: {e}", "error")
    
    def save_outputs(self, tailored_resume, cover_letter, job_title, company, job_description, match_score=0, match_summary=None):
        """Save tailored documents to output folder and show user where they are"""
        try:
            # Create timestamp
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            safe_company = company.replace(" ", "_").replace("/", "_")
            safe_title = job_title.replace(" ", "_").replace("/", "_")
            
            # Create base filename
            base_name = f"{safe_company}_{safe_title}_{timestamp}"
            
            # Ensure output directory exists
            OUTPUT_PATH.mkdir(parents=True, exist_ok=True)
            
            # Save tailored resume
            resume_path = OUTPUT_PATH / f"{base_name}_resume.txt"
            with open(resume_path, 'w', encoding='utf-8') as f:
                f.write(tailored_resume)
            
            # Save cover letter
            cover_letter_path = OUTPUT_PATH / f"{base_name}_cover_letter.txt"
            with open(cover_letter_path, 'w', encoding='utf-8') as f:
                f.write(cover_letter)
            
            # Save job description
            job_description_path = OUTPUT_PATH / f"{base_name}_job_description.txt"
            with open(job_description_path, 'w', encoding='utf-8') as f:
                f.write(job_description)
            
            # Add to database
            self.db_manager.add_application(
                job_title=job_title,
                company_name=company,
                job_url="",
                resume_path=str(resume_path),
                cover_letter_path=str(cover_letter_path),
                job_description_path=str(job_description_path),
                match_score=match_score,
                match_summary=json.dumps(match_summary) if match_summary else None
            )
            
            # SHOW USER WHERE FILES ARE SAVED (Fix #2)
            messagebox.showinfo(
                "Files Saved Successfully",
                f"Tailored documents saved to:\n\n{OUTPUT_PATH}\n\n"
                f"Resume: {base_name}_resume.txt\n"
                f"Cover Letter: {base_name}_cover_letter.txt\n"
                f"Job Description: {base_name}_job_description.txt"
            )
            
            self._log_message(f"Files saved: {base_name}_*.txt", "info")
        except Exception as e:
            self._log_message(f"Error saving files: {e}", "error")
            raise
    
    def clear_fields(self):
        """Clear all input fields"""
        self.job_title_entry.delete(0, tk.END)
        self.company_entry.delete(0, tk.END)
        self.job_desc_text.delete('1.0', tk.END)
        self.job_url_entry.delete(0, tk.END)
    
    def quit_application(self):
        """Quit the application with proper cleanup"""
        # Check if there are any ongoing operations
        if hasattr(self, '_processing') and self._processing:
            # Ask user if they want to force quit
            result = messagebox.askyesno(
                "Ongoing Process", 
                "A process is currently running. Are you sure you want to quit?\n\n"
                "WARNING: This may cause data loss or incomplete operations."
            )
            if not result:
                return
        
        # Close any open tooltip windows
        if hasattr(self, '_tooltip_window') and self._tooltip_window:
            try:
                self._tooltip_window.destroy()
            except:
                pass
        
        # Quit the application
        self.master.destroy()
    
    def set_ui_enabled(self, enabled):
        """Enable or disable UI elements during processing"""
        state = 'normal' if enabled else 'disabled'
        
        self.job_title_entry.config(state=state)
        self.company_entry.config(state=state)
        self.job_desc_text.config(state=state)
        self.job_url_entry.config(state=state)
        self.start_button.config(state=state)
        self.clear_button.config(state=state)
        self.upload_button.config(state=state)
        self.delete_button.config(state=state)
        self.set_active_button.config(state=state)
    
    def _show_tooltip(self, event, text):
        """Show tooltip with role definitions"""
        # Create tooltip window if it doesn't exist
        if not hasattr(self, '_tooltip_window') or not self._tooltip_window:
            self._tooltip_window = tk.Toplevel()
            self._tooltip_window.wm_overrideredirect(True)
            self._tooltip_window.configure(bg='lightyellow', relief='solid', borderwidth=1)
            
            # Create scrolled text widget for tooltip content to handle overflow
            tooltip_text = scrolledtext.ScrolledText(
                self._tooltip_window, 
                width=50, 
                height=15,
                wrap=tk.WORD,
                bg='lightyellow', 
                fg='black', 
                font=('Arial', 9),
                padx=5, 
                pady=3,
                bd=0,
                highlightthickness=0
            )
            tooltip_text.insert('1.0', text)
            tooltip_text.config(state='disabled')  # Make read-only
            tooltip_text.pack()
        
        # Position tooltip near mouse cursor but ensure it stays on screen
        x, y = event.x_root + 10, event.y_root + 10
        
        # Get screen dimensions
        screen_width = self.master.winfo_screenwidth()
        screen_height = self.master.winfo_screenheight()
        
        # Get tooltip dimensions
        self._tooltip_window.update_idletasks()  # Update to get accurate size
        tooltip_width = self._tooltip_window.winfo_reqwidth()
        tooltip_height = self._tooltip_window.winfo_reqheight()
        
        # Adjust position if tooltip would go off screen
        if x + tooltip_width > screen_width:
            x = screen_width - tooltip_width - 10
        if y + tooltip_height > screen_height:
            y = screen_height - tooltip_height - 10
        
        self._tooltip_window.wm_geometry(f"+{x}+{y}")
        self._tooltip_window.deiconify()
    
    def _hide_tooltip(self):
        """Hide tooltip window"""
        if hasattr(self, '_tooltip_window') and self._tooltip_window:
            self._tooltip_window.destroy()
            self._tooltip_window = None
    
    def _log_message(self, message, level='info'):
        """Add message to log window"""
        timestamp = datetime.now().strftime("%H:%M:%S")
        log_entry = f"[{timestamp}] {level.upper()}: {message}\n"
        
        # Insert at end
        self.log_text.insert(tk.END, log_entry)
        
        # Auto-scroll to bottom
        self.log_text.see(tk.END)
        
        # Also log to console
        logging.log(getattr(logging, level.upper()), message)
    
    def _check_queue(self):
        """Check for messages from worker threads"""
        try:
            while True:
                data = self.tailoring_queue.get_nowait()
                if data:
                    self.on_tailoring_complete(data)
        except queue.Empty:
            pass
        
        # Check again in 100ms
        self.master.after(100, self._check_queue)
    
    def _check_api_key(self):
        """Check if Gemini API key is configured with robust path detection"""
        # Try multiple possible .env locations
        possible_paths = [
            Path(__file__).parent.parent / ".env",  # Development path
            Path.cwd() / ".env",                    # Current working directory
            Path.home() / ".job_application_bot.env",  # User home directory
        ]
        
        api_key = None
        for env_path in possible_paths:
            if env_path.exists():
                load_dotenv(dotenv_path=env_path)
                api_key = os.getenv("GEMINI_API_KEY")
                if api_key and api_key != "your_api_key_here":
                    self._log_message(f"API key loaded from {env_path}", "info")
                    return  # Success, exit early
        
        # If we get here, no valid API key found
        messagebox.showwarning(
            "API Key Missing",
            "Gemini API key not found. Please set GEMINI_API_KEY in one of these locations:\n\n"
            "1. Project root: job-application-bot/.env\n"
            "2. Current directory: ./.env\n"
            "3. Home directory: ~/.job_application_bot.env\n\n"
            "You can get a free API key from: https://makersuite.google.com/app/apikey "
        )
        self._log_message("API key missing - please configure .env file", "warning")

def main():
    """Main entry point for the application"""
    print("DEBUG: Starting application...")
    root = tk.Tk()
    print("DEBUG: Tk root created")
    app = JobAppTkinter(root)
    print("DEBUG: JobAppTkinter instantiated")
    root.mainloop()
    print("DEBUG: Main loop ended")

if __name__ == "__main__":
    try:
        print("DEBUG: Entering main execution")
        main()
        print("DEBUG: Exited main execution normally")
    except Exception as e:
        import traceback
        print(f"DEBUG: Exception in main: {e}")
        traceback.print_exc()
        input("Press Enter to continue...")
